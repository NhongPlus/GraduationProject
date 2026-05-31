from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

def add_heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'

def add_image_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('[JPG: Chèn ảnh chụp màn hình tại đây]')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.add_run(caption)
    run2.bold = True
    run2.italic = True
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

fig = [0]
def next_fig(name):
    fig[0] += 1
    caption = f'Hình 4.{fig[0]}. {name}'
    add_image_placeholder(caption)

# ============================================================
# 4.3. SINH VIÊN
# ============================================================

add_heading2('4.3. Triển khai các chức năng dành cho sinh viên')

add_body(
    'Giao diện dành cho sinh viên được thiết kế tập trung vào trải nghiệm thi trực tuyến và theo dõi kết quả học tập. '
    'Sau khi đăng nhập, sinh viên có thể truy cập các chức năng thông qua thanh điều hướng bên trái, bao gồm: '
    'Trang chủ, Danh sách bài thi, Kết quả của tôi, Đánh giá học tập AI và Hồ sơ cá nhân.'
)

# 4.3.1
add_heading3('4.3.1. Giao diện đăng nhập và đổi mật khẩu lần đầu')

add_body(
    'Giao diện đăng nhập là điểm truy cập đầu tiên của hệ thống, được thiết kế theo bố cục chia đôi màn hình: '
    'bên trái là form đăng nhập với các trường email và mật khẩu, bên phải là carousel hình ảnh giới thiệu hệ thống. '
    'Sinh viên nhập thông tin đăng nhập và nhấn nút "Đăng nhập" để truy cập hệ thống. '
    'Trong trường hợp quên mật khẩu, sinh viên có thể mở modal "Quên mật khẩu" để gửi yêu cầu đặt lại qua email. '
    'Nếu phiên đăng nhập trước đó bị thu hồi do đăng nhập từ thiết bị khác, hệ thống hiển thị cảnh báo trên giao diện đăng nhập.'
)
next_fig('Giao diện đăng nhập')

add_body(
    'Khi sinh viên đăng nhập lần đầu tiên với tài khoản được cấp bởi quản trị viên, hệ thống bắt buộc đổi mật khẩu trước khi cho phép truy cập các chức năng khác. '
    'Giao diện đổi mật khẩu lần đầu hiển thị dưới dạng modal không thể đóng, yêu cầu nhập mật khẩu hiện tại, mật khẩu mới (tối thiểu 8 ký tự) và xác nhận mật khẩu mới. '
    'Sau khi đổi thành công, hệ thống tự động làm mới phiên đăng nhập và chuyển hướng đến trang chủ.'
)
next_fig('Giao diện đổi mật khẩu lần đầu')

# 4.3.2
add_heading3('4.3.2. Giao diện trang chủ sinh viên')

add_body(
    'Trang chủ sinh viên hiển thị bảng tổng quan về tình hình học tập và thi cử. '
    'Phần đầu là lời chào cá nhân hóa kèm theo các thẻ thống kê nhanh: số bài thi sắp tới, số bài đã hoàn thành và điểm trung bình. '
    'Bên dưới là bảng danh sách các bài thi sắp diễn ra với thông tin môn học, thời gian và trạng thái. '
    'Biểu đồ hiệu suất học tập so sánh điểm cá nhân với điểm trung bình của lớp theo từng bài thi, '
    'giúp sinh viên đánh giá vị trí của mình trong lớp. '
    'Phần kết quả gần đây liệt kê các bài thi đã hoàn thành với điểm số và trạng thái đạt/không đạt.'
)
next_fig('Giao diện trang chủ sinh viên')

# 4.3.3
add_heading3('4.3.3. Giao diện danh sách bài thi')

add_body(
    'Giao diện danh sách bài thi hiển thị tất cả các bài thi mà sinh viên có thể tham gia. '
    'Phần đầu trang hiển thị các thẻ thống kê tổng quan: tổng số bài thi, số bài đang mở, số bài đã hoàn thành. '
    'Bảng danh sách bài thi bao gồm các cột: tên bài thi, môn học, thời lượng, số lần thi còn lại, thời gian mở/đóng và trạng thái. '
    'Sinh viên có thể tìm kiếm bài thi theo tên hoặc môn học. '
    'Với mỗi bài thi, sinh viên có thể nhấn "Vào thi" để bắt đầu làm bài hoặc "Xem kết quả" nếu đã hoàn thành. '
    'Các bài thi chưa mở hoặc đã hết hạn được đánh dấu trạng thái tương ứng và không cho phép truy cập.'
)
next_fig('Giao diện danh sách bài thi của sinh viên')

# 4.3.4
add_heading3('4.3.4. Giao diện làm bài thi trực tuyến')

add_body(
    'Giao diện làm bài thi là chức năng cốt lõi của hệ thống, được thiết kế toàn màn hình (fullscreen) '
    'với bố cục gồm ba phần chính: thanh tiêu đề, khu vực hiển thị câu hỏi và bảng điều hướng câu hỏi.'
)

add_body(
    'Thanh tiêu đề hiển thị tên bài thi, mã đề (D01, D02,...) và đồng hồ đếm ngược thời gian còn lại. '
    'Khu vực câu hỏi hiển thị nội dung câu hỏi (hỗ trợ rich text, hình ảnh, âm thanh) cùng với các lựa chọn trả lời '
    'tùy theo loại câu hỏi: radio button cho trắc nghiệm, ô nhập văn bản cho tự luận và điền khuyết. '
    'Bảng điều hướng bên phải hiển thị danh sách câu hỏi dưới dạng các nút đánh số, phân biệt bằng màu sắc '
    'giữa câu đã trả lời, chưa trả lời và câu đã đánh dấu (flag). '
    'Sinh viên có thể di chuyển giữa các câu hỏi bằng nút Trước/Sau hoặc nhấn trực tiếp vào số câu.'
)
next_fig('Giao diện làm bài thi trực tuyến')

add_body(
    'Hệ thống tích hợp cơ chế tự động lưu bài (autosave) định kỳ gửi câu trả lời lên server, '
    'đảm bảo không mất dữ liệu khi gặp sự cố mạng. Khi hết thời gian, bài thi được tự động nộp. '
    'Sinh viên cũng có thể chủ động nhấn nút "Nộp bài" kèm hộp thoại xác nhận hiển thị '
    'số câu đã trả lời và số câu còn trống.'
)

add_body(
    'Về chống gian lận, hệ thống áp dụng nhiều biện pháp bảo vệ tính toàn vẹn bài thi: '
    'bắt buộc chế độ toàn màn hình, phát hiện chuyển tab hoặc mất focus cửa sổ, '
    'chặn chuột phải và thao tác sao chép/dán, ghi nhận sự kiện vi phạm và gửi lên server theo thời gian thực. '
    'Hệ thống áp dụng cơ chế "strike" — sau một số lần vi phạm nhất định, bài thi sẽ bị khóa '
    'và tự động nộp kèm đếm ngược cảnh báo. Kết nối Socket.IO cho phép giảng viên gửi cảnh báo '
    'hoặc ép nộp bài từ xa trong quá trình thi.'
)
next_fig('Giao diện cảnh báo vi phạm khi làm bài thi')

# 4.3.5
add_heading3('4.3.5. Giao diện kết quả bài thi')

add_body(
    'Sau khi nộp bài, sinh viên được chuyển đến giao diện kết quả bài thi hiển thị thông tin chi tiết về bài làm. '
    'Phần đầu là các thẻ tóm tắt: tổng điểm, số câu đúng/sai, thời gian làm bài và trạng thái đạt/không đạt. '
    'Nếu bài thi có câu tự luận chưa được chấm, hệ thống hiển thị trạng thái "Đang chờ chấm" hoặc "Chấm một phần".'
)

add_body(
    'Phần đánh giá học tập từ AI phân tích kết quả theo từng chương/chủ đề, hiển thị thanh tiến trình '
    'tỷ lệ đúng và đưa ra nhận xét về điểm mạnh, điểm yếu cùng gợi ý cải thiện cho sinh viên. '
    'Bên dưới là danh sách chi tiết từng câu hỏi dưới dạng accordion có thể mở rộng, '
    'hiển thị nội dung câu hỏi, câu trả lời của sinh viên, đáp án đúng (với câu khách quan), '
    'điểm đạt được và nhận xét của giảng viên (nếu có).'
)
next_fig('Giao diện kết quả bài thi')

# 4.3.6
add_heading3('4.3.6. Giao diện lịch sử các lần thi')

add_body(
    'Giao diện "Kết quả của tôi" tổng hợp toàn bộ lịch sử thi của sinh viên. '
    'Phần thống kê tổng quan hiển thị tổng số lần thi, số bài đang chờ chấm thủ công và điểm trung bình. '
    'Bảng danh sách liệt kê các phiên thi đã thực hiện với thông tin: tên bài thi, mã đề, '
    'thời gian bắt đầu/kết thúc, điểm số và trạng thái chấm bài. '
    'Sinh viên có thể nhấn vào từng dòng để xem chi tiết kết quả bài thi tương ứng.'
)
next_fig('Giao diện lịch sử các lần thi')

# 4.3.7
add_heading3('4.3.7. Giao diện đánh giá học tập từ AI')

add_body(
    'Chức năng đánh giá học tập từ AI sử dụng MiniMax API để phân tích toàn bộ dữ liệu học tập của sinh viên '
    'và đưa ra dự đoán kết quả cho môn học mục tiêu. Giao diện gồm ba phần chính.'
)

add_body(
    'Phần đầu là bảng lịch sử điểm hiển thị kết quả các môn đã học, làm dữ liệu đầu vào cho mô hình AI. '
    'Phần thứ hai là bộ chọn môn học mục tiêu theo cấu trúc danh mục (nhóm môn → môn), '
    'hệ thống kiểm tra điều kiện tiên quyết và số môn tối thiểu đã hoàn thành trước khi cho phép dự đoán. '
    'Phần thứ ba hiển thị kết quả dự đoán bao gồm: điểm dự kiến, xếp loại (A/B/C/D/F), '
    'phân tích điểm mạnh và điểm yếu, các chương cần ôn tập, '
    'tổng hợp câu trả lời sai và lời khuyên học tập cụ thể từ AI.'
)
next_fig('Giao diện đánh giá học tập từ AI')
next_fig('Giao diện kết quả dự đoán điểm từ AI')

# 4.3.8
add_heading3('4.3.8. Giao diện hồ sơ cá nhân')

add_body(
    'Giao diện hồ sơ cá nhân cho phép sinh viên xem thông tin tài khoản và quản lý bảo mật. '
    'Tab "Hồ sơ" hiển thị họ tên, vai trò (badge), email và các thông tin cá nhân. '
    'Tab "Cài đặt" cung cấp chức năng đổi mật khẩu với các trường: mật khẩu hiện tại, '
    'mật khẩu mới và xác nhận mật khẩu mới. '
    'Ngoài ra, sinh viên có thể gửi yêu cầu đặt lại mật khẩu đến quản trị viên '
    'trong trường hợp không nhớ mật khẩu hiện tại, và xem lịch sử các yêu cầu đã gửi. '
    'Nút đăng xuất cho phép kết thúc phiên làm việc và quay về trang đăng nhập.'
)
next_fig('Giao diện hồ sơ cá nhân')

# ============================================================
# 4.4. GIẢNG VIÊN
# ============================================================

add_heading2('4.4. Triển khai các chức năng dành cho giảng viên')

add_body(
    'Giao diện dành cho giảng viên mở rộng các chức năng quản lý đề thi, chấm điểm và theo dõi kết quả học tập. '
    'Sau khi đăng nhập, giảng viên truy cập hệ thống qua thanh điều hướng bên trái với các mục: '
    'Trang chủ, Ngân hàng câu hỏi, Bài thi, Chấm bài, Phân tích điểm và Quản lý sinh viên.'
)

# 4.4.1
add_heading3('4.4.1. Giao diện trang chủ giảng viên')

add_body(
    'Trang chủ giảng viên hiển thị bảng tổng quan hoạt động giảng dạy và quản lý thi cử. '
    'Phần thống kê nhanh bao gồm các thẻ số liệu: tổng số sinh viên quản lý, số bài thi đã tạo, '
    'số phiên thi đang diễn ra và số bài chờ chấm điểm.'
)

add_body(
    'Bên dưới là bảng danh sách sinh viên với chức năng tìm kiếm, lọc theo lớp hành chính và phân trang. '
    'Phần hoạt động gần đây hiển thị các sự kiện mới nhất trong hệ thống như bài thi mới được nộp, '
    'phiên thi được mở, kết quả chấm bài, có thể lọc theo trạng thái, thời gian và từ khóa. '
    'Giảng viên có thể nhanh chóng chuyển đến quản lý sinh viên hoặc các chức năng khác từ trang chủ.'
)
next_fig('Giao diện trang chủ giảng viên')

# 4.4.2
add_heading3('4.4.2. Giao diện quản lý đề thi')

add_body(
    'Giảng viên có thể quản lý đề thi thông qua danh sách bài thi và trang soạn đề. '
    'Tại danh sách bài thi, giao diện của giảng viên có thêm cột "Phiên đang thi" hiển thị số sinh viên đang làm bài, '
    'nút "Mở phiên thi" để kích hoạt bài thi và "Ép nộp bài" để cưỡng chế nộp toàn bộ. '
    'Menu hành động cho phép sửa câu hỏi, đặt giờ, quản lý phiên thi và xóa đề thi.'
)
next_fig('Giao diện danh sách bài thi của giảng viên')

add_body(
    'Trang soạn đề thi cho phép tạo mới hoặc chỉnh sửa đề thi. '
    'Giảng viên nhập thông tin gồm tiêu đề, mô tả, thời lượng, lịch trình mở/đóng, '
    'lớp hành chính, môn học và số lượng mã đề (tối đa 4 mã: D01 đến D04). '
    'Phần quản lý câu hỏi theo từng mã đề cho phép thêm câu hỏi thủ công, '
    'upload hình ảnh hoặc chọn câu từ ngân hàng câu hỏi.'
)
next_fig('Giao diện soạn đề thi — phần thông tin chung')
next_fig('Giao diện soạn đề thi — phần quản lý câu hỏi')

# 4.4.3
add_heading3('4.4.3. Giao diện quản lý câu hỏi và ngân hàng câu hỏi')

add_body(
    'Ngân hàng câu hỏi là nơi giảng viên lưu trữ và quản lý câu hỏi độc lập với đề thi, '
    'giúp tái sử dụng cho nhiều bài thi khác nhau. '
    'Giao diện hiển thị bảng danh sách câu hỏi với thông tin nội dung, loại câu hỏi (trắc nghiệm/tự luận), '
    'độ khó (Dễ, Trung bình, Khó) được đánh dấu bằng badge màu, chương, môn học và số lần sử dụng. '
    'Giảng viên có thể tìm kiếm, lọc theo nhiều tiêu chí (loại câu hỏi, độ khó, môn học, chương), '
    'tạo/sửa/xóa câu hỏi đơn lẻ hoặc xóa hàng loạt bằng checkbox chọn nhiều.'
)
next_fig('Giao diện ngân hàng câu hỏi')

# 4.4.4
add_heading3('4.4.4. Giao diện import đề thi từ file Word')

add_body(
    'Hệ thống hỗ trợ import câu hỏi từ file Word (.docx) giúp giảng viên tạo đề thi nhanh chóng từ tài liệu có sẵn. '
    'Quy trình gồm: tải template mẫu, upload file Word (kéo thả hoặc chọn file), '
    'xem trước danh sách câu hỏi được trích xuất trong modal preview, và xác nhận import. '
    'Ngoài ra, hệ thống còn hỗ trợ chức năng AI Recompose sử dụng MiniMax API '
    'để tự động phân tích và tái cấu trúc nội dung file Word thành câu hỏi chuẩn. '
    'Giảng viên cũng có thể upload file media đi kèm (hình ảnh trong file ZIP) '
    'để gắn với các câu hỏi có minh họa.'
)
next_fig('Giao diện import đề thi từ file Word')

# 4.4.5
add_heading3('4.4.5. Giao diện theo dõi phiên thi')

add_body(
    'Giao diện theo dõi phiên thi cho phép giảng viên quản lý các phiên thi của sinh viên cho một bài thi cụ thể. '
    'Bảng danh sách hiển thị thông tin sinh viên, mã đề, thời gian bắt đầu/kết thúc, điểm số và trạng thái chấm bài. '
    'Phần thống kê tổng quan cho biết tổng số phiên, số phiên đang thi, đã nộp và đã chấm. '
    'Giảng viên có thể chuyển đến giao diện chấm bài hoặc cấp quyền thi lại cho sinh viên '
    'thông qua modal nhập lý do thi lại, đồng thời có thể thu hồi quyền thi lại đã cấp.'
)
next_fig('Giao diện theo dõi phiên thi')

# 4.4.6
add_heading3('4.4.6. Giao diện chấm điểm bài thi')

add_body(
    'Hệ thống chấm điểm gồm hai phần: danh sách bài chờ chấm và giao diện chấm chi tiết. '
    'Tại danh sách chờ chấm, hệ thống tự động quét các bài thi có câu tự luận '
    'và liệt kê các phiên cần chấm thủ công, có thể lọc theo từ khóa, môn học và khoảng thời gian.'
)
next_fig('Giao diện danh sách bài chờ chấm')

add_body(
    'Tại giao diện chấm chi tiết, phần trên hiển thị thông tin sinh viên và bảng tóm tắt các câu trắc nghiệm '
    'đã được hệ thống chấm tự động. Phần dưới là khu vực chấm câu tự luận — '
    'với mỗi câu, giảng viên đọc bài làm của sinh viên, nhập điểm số và viết nhận xét, '
    'sau đó nhấn "Lưu kết quả" để hoàn tất việc chấm.'
)
next_fig('Giao diện chấm điểm chi tiết')

# 4.4.7
add_heading3('4.4.7. Giao diện quản lý danh sách sinh viên')

add_body(
    'Giảng viên có thể quản lý danh sách sinh viên thuộc quyền quản lý thông qua giao diện quản lý sinh viên. '
    'Bảng danh sách hiển thị thông tin: họ tên, username, email, lớp hành chính, trạng thái tài khoản. '
    'Giảng viên có thể tìm kiếm sinh viên, chỉnh sửa thông tin cá nhân và xóa tài khoản sinh viên. '
    'Hệ thống hỗ trợ hiển thị/ẩn mật khẩu tạm thời của sinh viên mới được tạo. '
    'Giảng viên cũng có thể gửi email bảng điểm hàng loạt cho toàn bộ hoặc một nhóm sinh viên được chọn, '
    'và xem bảng điểm tổng (transcript) của từng sinh viên.'
)
next_fig('Giao diện quản lý danh sách sinh viên')

# 4.4.8
add_heading3('4.4.8. Giao diện bảng điểm và transcript sinh viên')

add_body(
    'Giảng viên có thể xem bảng điểm tổng hợp của từng sinh viên thông qua modal transcript. '
    'Bảng điểm hiển thị thông tin cá nhân sinh viên, danh sách các bài thi đã hoàn thành '
    'với điểm số, xếp loại và thời gian thi. Phần tổng hợp hiển thị GPA theo thang điểm 4 và thang điểm 10. '
    'Giảng viên có thể in bảng điểm trực tiếp, xuất file HTML hoặc CSV, '
    'và gửi bảng điểm qua email cho sinh viên.'
)
next_fig('Giao diện bảng điểm tổng hợp (transcript) sinh viên')

# 4.4.9
add_heading3('4.4.9. Giao diện thống kê kết quả học tập')

add_body(
    'Giảng viên có thể xem thống kê kết quả học tập thông qua giao diện phân tích điểm số. '
    'Sau khi chọn lớp hành chính và môn học từ bộ lọc, giao diện hiển thị các thẻ thống kê tổng quan '
    'bao gồm: điểm trung bình, tỷ lệ đạt, điểm cao nhất và điểm thấp nhất.'
)

add_body(
    'Biểu đồ cột phân bố điểm (histogram) trực quan hóa sự phân bố điểm của sinh viên theo các khoảng điểm. '
    'Bảng chi tiết bên dưới liệt kê kết quả từng bài thi với thông tin số phiên, điểm trung bình, '
    'tỷ lệ đạt và điểm cao nhất/thấp nhất cho mỗi bài.'
)
next_fig('Giao diện thống kê kết quả học tập')

add_body(
    'Ngoài ra, tab bảng điểm trong giao diện quản lý sinh viên cho phép giảng viên chọn bài thi, '
    'tìm kiếm sinh viên và xem bảng điểm dạng bảng. Giảng viên có thể xuất bảng điểm ra file CSV '
    'và gửi email thông báo điểm cho toàn bộ hoặc nhóm sinh viên được chọn.'
)
next_fig('Giao diện bảng điểm theo bài thi')

# ============================================================
# 4.5. QUẢN TRỊ VIÊN
# ============================================================

add_heading2('4.5. Triển khai các chức năng dành cho quản trị viên')

add_body(
    'Hệ thống quản trị được thiết kế tách biệt với giao diện sinh viên và giảng viên '
    'nhằm đảm bảo tính rõ ràng trong quá trình quản lý. '
    'Sau khi đăng nhập, quản trị viên truy cập các mục chức năng thông qua thanh menu bên trái, bao gồm: '
    'Trang chủ, Quản lý chương trình đào tạo, Quản lý môn học, Phân tích điểm, '
    'Quản lý người dùng, Quản lý lớp hành chính và các công cụ quản trị.'
)

# 4.5.1
add_heading3('4.5.1. Giao diện trang chủ quản trị viên')

add_body(
    'Trang chủ quản trị viên hiển thị bảng tổng quan toàn hệ thống với các thẻ số liệu: '
    'tổng số tài khoản, số bài thi đã tạo, số phiên thi và số lớp hành chính. '
    'Bên dưới là bảng danh sách sinh viên có thể tìm kiếm và lọc theo lớp, '
    'cùng với danh sách hoạt động gần đây có thể lọc theo trạng thái, thời gian và từ khóa. '
    'Quản trị viên có thể nhanh chóng chuyển đến trang quản lý sinh viên từ trang chủ.'
)
next_fig('Giao diện trang chủ quản trị viên')

# 4.5.2
add_heading3('4.5.2. Giao diện quản lý người dùng')

add_body(
    'Giao diện quản lý người dùng cho phép quản trị viên quản lý toàn bộ tài khoản trong hệ thống '
    'bao gồm sinh viên và giảng viên. '
    'Bảng danh sách hiển thị thông tin: họ tên, username, email, vai trò, lớp hành chính và trạng thái. '
    'Hệ thống cung cấp hai ô tìm kiếm riêng biệt cho sinh viên và giảng viên, '
    'bộ lọc theo lớp hành chính và phân trang.'
)

add_body(
    'Quản trị viên có thể tạo tài khoản mới (chọn vai trò sinh viên hoặc giảng viên), '
    'chỉnh sửa thông tin người dùng, kích hoạt/vô hiệu hóa tài khoản, '
    'đặt lại mật khẩu (hiển thị mật khẩu tạm thời) và xóa tài khoản đơn lẻ hoặc hàng loạt. '
    'Chức năng hiển thị/ẩn mật khẩu tạm thời giúp quản trị viên thông báo cho người dùng mới.'
)
next_fig('Giao diện quản lý người dùng')
next_fig('Giao diện tạo tài khoản người dùng mới')

# 4.5.3
add_heading3('4.5.3. Giao diện quản lý lớp hành chính')

add_body(
    'Giao diện quản lý lớp hành chính gồm hai phần: danh sách lớp và chi tiết lớp. '
    'Phần danh sách hiển thị các lớp hành chính với thông tin: tên lớp, chương trình đào tạo, '
    'khóa nhập học, giảng viên quản lý, sĩ số hiện tại và sĩ số dự kiến. '
    'Quản trị viên có thể tạo lớp mới bằng cách chọn chương trình đào tạo, khóa, phân ban, '
    'tên hiển thị và giảng viên chủ nhiệm.'
)
next_fig('Giao diện danh sách lớp hành chính')

add_body(
    'Khi chọn một lớp, giao diện chi tiết hiển thị danh sách sinh viên của lớp đó. '
    'Quản trị viên có thể thêm sinh viên bằng nhiều cách: chọn từ danh sách sinh viên chưa có lớp, '
    'chuyển sinh viên từ lớp khác (cho phép chuyển lớp), thêm thủ công bằng form nhập thông tin, '
    'hoặc import hàng loạt từ file Excel. '
    'Quy trình import Excel gồm: tải template mẫu, upload file, xem trước danh sách trong modal preview '
    'và xác nhận import. Quản trị viên cũng có thể xóa sinh viên khỏi lớp.'
)
next_fig('Giao diện chi tiết lớp hành chính')
next_fig('Giao diện import sinh viên từ file Excel')

# 4.5.4
add_heading3('4.5.4. Giao diện quản lý chương trình đào tạo')

add_body(
    'Giao diện quản lý chương trình đào tạo cho phép quản trị viên quản lý các ngành học trong hệ thống. '
    'Bảng danh sách hiển thị: mã chương trình, tên chương trình, mô tả, trạng thái hoạt động '
    'và danh sách giảng viên phụ trách. '
    'Quản trị viên có thể tạo mới, chỉnh sửa, xóa chương trình và bật/tắt trạng thái hoạt động.'
)

add_body(
    'Chức năng phân công giảng viên cho phép quản trị viên mở modal chọn nhiều giảng viên '
    'từ danh sách có sẵn và gán vào chương trình đào tạo. '
    'Giảng viên được phân công sẽ có quyền truy cập các môn học và lớp thuộc chương trình tương ứng.'
)
next_fig('Giao diện quản lý chương trình đào tạo')

# 4.5.5
add_heading3('4.5.5. Giao diện quản lý môn học và nhóm môn')

add_body(
    'Giao diện quản lý môn học và nhóm môn là module quản lý danh mục học phần phức tạp nhất của hệ thống. '
    'Sau khi chọn chương trình đào tạo từ dropdown, giao diện hiển thị hai tab: '
    '"Nhóm môn" và "Môn học".'
)

add_body(
    'Tab "Nhóm môn" cho phép quản trị viên tạo, sửa, xóa các nhóm môn học trong chương trình '
    '(ví dụ: Đại cương, Cơ sở ngành, Chuyên ngành). '
    'Chức năng "Gán từ kho" cho phép sao chép nhóm môn và môn học từ kho dữ liệu gốc vào chương trình, '
    'giúp thiết lập nhanh danh mục cho chương trình mới.'
)
next_fig('Giao diện quản lý nhóm môn học')

add_body(
    'Tab "Môn học" hiển thị bảng danh sách môn học với thông tin: mã môn, tên môn, số tín chỉ, '
    'học kỳ, phân loại, nhóm môn, môn tiên quyết và trạng thái. '
    'Quản trị viên có thể tạo/sửa/xóa môn học, thiết lập môn tiên quyết, bật/tắt trạng thái '
    'và xóa hàng loạt bằng checkbox. '
    'Hệ thống hỗ trợ import môn học từ file Excel với quy trình: tải template, upload, preview và xác nhận.'
)
next_fig('Giao diện quản lý môn học')

# 4.5.6
add_heading3('4.5.6. Giao diện quản lý yêu cầu đặt lại mật khẩu')

add_body(
    'Khi sinh viên hoặc giảng viên gửi yêu cầu đặt lại mật khẩu, quản trị viên nhận được '
    'danh sách yêu cầu chờ xử lý tại giao diện này. '
    'Bảng danh sách hiển thị: tên người yêu cầu, email, thời gian gửi và trạng thái. '
    'Với mỗi yêu cầu, quản trị viên có thể phê duyệt (hệ thống tự động tạo mật khẩu tạm thời '
    'và hiển thị để thông báo cho người dùng) hoặc từ chối kèm ghi chú lý do. '
    'Quản trị viên cũng có thể làm mới danh sách để kiểm tra yêu cầu mới.'
)
next_fig('Giao diện quản lý yêu cầu đặt lại mật khẩu')

# 4.5.7
add_heading3('4.5.7. Giao diện nhật ký hoạt động')

add_body(
    'Giao diện nhật ký hoạt động (Audit Log) ghi nhận toàn bộ các hành động quan trọng trong hệ thống '
    'nhằm phục vụ mục đích giám sát và truy vết. '
    'Bảng nhật ký hiển thị: thời gian, người thực hiện, loại hành động, mô tả chi tiết và đối tượng bị tác động. '
    'Các loại hành động được ghi nhận bao gồm: đăng nhập/đăng xuất, tạo/sửa/xóa tài khoản, '
    'nộp bài thi, chấm điểm, đặt lại mật khẩu và các thao tác quản trị khác. '
    'Quản trị viên có thể lọc nhật ký theo loại hành động và phân trang để duyệt lịch sử.'
)
next_fig('Giao diện nhật ký hoạt động')

# 4.5.8
add_heading3('4.5.8. Giao diện báo cáo hệ thống')

add_body(
    'Giao diện báo cáo hệ thống cung cấp cái nhìn tổng quan về toàn bộ hoạt động của hệ thống '
    'thông qua các chỉ số KPI quan trọng. Phần thống kê tài khoản hiển thị tổng số tài khoản, '
    'số sinh viên, giảng viên và quản trị viên.'
)

add_body(
    'Phần thống kê thi cử bao gồm: tổng số bài thi, tổng phiên thi, tỷ lệ hoàn thành, '
    'tỷ lệ đạt, điểm trung bình toàn hệ thống. '
    'Phần thống kê toàn vẹn bài thi hiển thị số vi phạm trong 24 giờ qua, '
    'số phiên bị đánh dấu và số bài chờ chấm điểm thủ công. '
    'Bên dưới là bảng danh sách các bài thi gần đây với thông tin chi tiết về số phiên, '
    'trạng thái và kết quả tổng hợp. '
    'Các thanh tiến trình (progress bar) trực quan hóa tỷ lệ hoàn thành và tỷ lệ đạt '
    'giúp quản trị viên nhanh chóng đánh giá hiệu quả tổng thể của hệ thống.'
)
next_fig('Giao diện báo cáo hệ thống')

output_path = r'c:\VS-Code\GraduationProject\docs\BaoCao_Chuong4_GiaoDien.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
print(f'Total figures: {fig[0]}')
