# -*- coding: utf-8 -*-
"""
Sinh báo cáo tiến độ tuần (Word) cho đồ án Hệ thống thi trực tuyến.

Chạy: python scripts/generate_bao_cao_tuan.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor

DE_TAI = "XÂY DỰNG HỆ THỐNG THI TRỰC TUYẾN\n(ONLINE EXAMINATION SYSTEM)"
TUAN_BAT_DAU = date(2026, 5, 11)
TUAN_KET_THUC = date(2026, 5, 18)
SO_COMMIT = 28


def set_body_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)


def fmt_para(p, justify: bool = True) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)


def add_center(doc: Document, text: str, *, bold: bool = True, size: int = 13) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    fmt_para(p, justify=False)


def add_p(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    fmt_para(p)


def add_h(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h.paragraph_format.line_spacing = 1.5


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    fmt_para(p)


def add_table_row(table, cells: list[str], *, bold: bool = False) -> None:
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = bold


def build_report() -> Document:
    doc = Document()
    set_body_style(doc)

    # --- Bìa ---
    add_center(doc, "BÁO CÁO TIẾN ĐỘ TUẦN", size=16)
    add_center(doc, "ĐỒ ÁN TỐT NGHIỆP", size=14)
    doc.add_paragraph()
    add_center(doc, DE_TAI, size=13)
    doc.add_paragraph()
    add_center(
        doc,
        f"Tuần: {TUAN_BAT_DAU.strftime('%d/%m/%Y')} – {TUAN_KET_THUC.strftime('%d/%m/%Y')}",
        bold=False,
    )
    add_center(doc, f"Ngày lập báo cáo: {date.today().strftime('%d/%m/%Y')}", bold=False)
    doc.add_page_break()

    # --- Tóm tắt ---
    add_h(doc, "1. Tóm tắt tuần", 1)
    add_p(
        doc,
        f"Trong khoảng thời gian từ {TUAN_BAT_DAU.strftime('%d/%m/%Y')} đến "
        f"{TUAN_KET_THUC.strftime('%d/%m/%Y')}, dự án Hệ thống thi trực tuyến có "
        f"{SO_COMMIT} commit trên nhánh chính, tập trung vào ba hướng lớn: (1) hoàn thiện "
        "nền tảng API, cơ sở dữ liệu và giao diện quản trị; (2) triển khai luồng thi trực "
        "tuyến đa mã đề với Socket.IO; (3) sửa lỗi chấm điểm trắc nghiệm, xem lại bài làm "
        "và tăng cường chống gian lận khi sinh viên làm bài."
    )
    add_p(
        doc,
        "Tuần này đánh dấu bước chuyển từ giai đoạn xây dựng module rời rạc sang giai đoạn "
        "tích hợp end-to-end: giáo viên soạn đề nhiều mã đề, bắt đầu ca thi realtime, sinh viên "
        "làm bài fullscreen có autosave, hệ thống chấm TN tự động và hiển thị kết quả/review "
        "sau khi nộp bài. Đồng thời đã cấu hình triển khai production trỏ API "
        "api.nhongplus.id.vn thay vì localhost."
    )

    # --- Thống kê ---
    add_h(doc, "2. Thống kê công việc", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Chỉ tiêu"
    hdr[1].text = "Giá trị"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
    rows = [
        ("Số commit (git)", str(SO_COMMIT)),
        ("Ngày có commit đầu tiên trong tuần", "13/05/2026"),
        ("Ngày commit mới nhất", "18/05/2026"),
        ("Migration DB mới", "010–024 (password reset, audit, exam versions, question bank, runtime, proctoring, autosave…)"),
        ("File Word mẫu đề thi", "exam_template_GiaoVien.docx"),
        ("Tài liệu tổng hợp", "DO_AN_MASTER.md"),
    ]
    for a, b in rows:
        add_table_row(table, [a, b])
    doc.add_paragraph()

    # --- Chi tiết theo nhóm ---
    add_h(doc, "3. Các thay đổi chính theo nhóm chức năng", 1)

    add_h(doc, "3.1. Ngân hàng câu hỏi & import đề Word", 2)
    add_bullet(doc, "Trang Question Bank: chọn môn học, import hàng loạt từ file .docx.")
    add_bullet(doc, "Modal ExamImportPreviewModal: xem trước câu hỏi sau khi parse Word trước khi lưu.")
    add_bullet(doc, "Backend: cải thiện examImport.service (Mammoth), hỗ trợ tag [LOAI:TN]/[LOAI:TL], file mẫu exam_template_GiaoVien.docx.")
    add_bullet(doc, "ExamAuthoring: làm rõ import Word theo từng mã đề; sao chép câu hỏi giữa các mã đề.")

    add_h(doc, "3.2. API Backend, cơ sở dữ liệu & tài liệu", 2)
    add_bullet(doc, "Thêm nhiều migration: exam_versions, question_bank, exam_sharing, exam_runtime_state, audit_logs, user_notifications, password_reset, exam_collaborators, admin_classes, proctoring, autosave timestamps…")
    add_bullet(doc, "API mới/bổ sung: question bank, exam sharing, shuffle câu hỏi, score analytics, export kết quả, audit log, system report, password reset admin, notification unread-count.")
    add_bullet(doc, "OpenAPI (openapi.yaml) và ROLES_AND_PERMISSIONS.md; script seed-data, check-migrations.")
    add_bullet(doc, "Tự chạy migration khi server khởi động (runMigrations.ts) — giảm lỗi thiếu bảng trên production.")
    add_bullet(doc, "Email service (Nodemailer) và tài liệu EMAIL_SETUP.md; quên mật khẩu qua token.")

    add_h(doc, "3.3. Thi trực tuyến đa mã đề & phiên thi realtime", 2)
    add_bullet(doc, "Hỗ trợ nhiều mã đề (num_versions), xáo trộn câu hỏi/đáp án, lưu option_map và question_order.")
    add_bullet(doc, "Socket.IO: giáo viên bắt đầu ca thi, cảnh báo, ép nộp bài; sinh viên đồng bộ trạng thái.")
    add_bullet(doc, "ExamSessions, ExamRuntimeState; cộng tác giám khảo (exam collaborators, proctoring dashboard).")
    add_bullet(doc, "Trang ExamList: hiển thị trạng thái Đã làm, Ép nộp bài; sửa lỗi redirect sai khi bấm Làm bài.")

    add_h(doc, "3.4. Làm bài, chấm điểm & xem kết quả", 2)
    add_bullet(doc, "ExamTake: bắt buộc fullscreen khi GV bắt đầu thi; autosave; redirect về trang chính 3 giây sau nộp bài.")
    add_bullet(doc, "Chấm TN: chuẩn hóa so khớp đáp án (chữ cái hoặc nội dung), chấm qua display key và option_map sau xáo trộn.")
    add_bullet(doc, "ExamResult: hiển thị lại đáp án đúng, option đã xáo trộn; sửa parseJson JSONB từ PostgreSQL.")
    add_bullet(doc, "Bổ sung unit test: examMcqGrading, examSessionReview, useExamTakeState, proctoringSocket…")
    add_bullet(doc, "Trang Grading: danh sách bài cần chấm tự luận.")

    add_h(doc, "3.5. Chống gian lận (integrity & proctoring)", 2)
    add_bullet(doc, "Báo cáo vi phạm (rời tab, thoát fullscreen…) lên server qua REST và Socket.")
    add_bullet(doc, "Migration proctoring_enhancements; model examProctor, examIntegrity.")
    add_bullet(doc, "ProctoringDashboard: theo dõi phiên thi realtime.")

    add_h(doc, "3.6. Giao diện quản trị & trải nghiệm", 2)
    add_bullet(doc, "Trang Admin: AuditLog, PasswordResetManagement, SubjectManagement, SystemReport.")
    add_bullet(doc, "ScoreAnalytics, Prediction (AI MiniMax), cải thiện Navbar và i18n (vi/en/ja).")
    add_bullet(doc, "NotificationBell: xử lý lỗi rõ ràng khi thiếu bảng/route thông báo.")

    add_h(doc, "3.7. Triển khai & cấu hình môi trường", 2)
    add_bullet(doc, "Tách .env.development / .env.production; VITE_API_URL trỏ api.nhongplus.id.vn trên production.")
    add_bullet(doc, "app.config.ts: không dùng localhost khi deploy domain công khai.")
    add_bullet(doc, "Gộp tài liệu đồ án vào DO_AN_MASTER.md; xóa các file markdown trùng lặp.")

    # --- Timeline ---
    add_h(doc, "4. Dòng thời gian commit (tóm tắt)", 1)
    timeline = [
        ("13/05", "Ngân hàng câu hỏi + import Word; đợt lớn API/migrations/FE; cấu hình production API; DO_AN_MASTER.md."),
        ("17/05", "Đa mã đề & phiên thi online; hàng loạt fix ExamList/ExamTake/MCQ grading/review; auto migration; notification."),
        ("18/05", "Báo cáo vi phạm integrity lên server; cải thiện luồng ExamTake."),
    ]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Ngày"
    t2.rows[0].cells[1].text = "Nội dung"
    for c in t2.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
    for ngay, nd in timeline:
        add_table_row(t2, [ngay, nd])
    doc.add_paragraph()

    # --- File quan trọng ---
    add_h(doc, "5. Tệp và thư mục đáng chú ý", 1)
    files = [
        "BackEnd/server/src/services/exam.service.ts — logic thi, chấm, review",
        "BackEnd/server/src/utils/examMcqGrading.ts — chấm trắc nghiệm",
        "BackEnd/server/src/socket/examSocket.ts — realtime ca thi",
        "FrontEnd/client/src/pages/main/Exam/ExamTake.tsx — giao diện làm bài",
        "FrontEnd/client/src/pages/main/Exam/ExamResult.tsx — kết quả & review",
        "FrontEnd/client/src/pages/main/Exam/QuestionBank.tsx — ngân hàng câu hỏi",
        "DO_AN_MASTER.md — tài liệu tổng hợp đồ án",
        "BackEnd/server/openapi.yaml — mô tả API",
    ]
    for f in files:
        add_bullet(doc, f)

    # --- Khó khăn & hướng tiếp ---
    add_h(doc, "6. Khó khăn đã gặp và hướng tuần tới", 1)
    add_p(
        doc,
        "Khó khăn chính trong tuần liên quan đến chấm trắc nghiệm sau khi xáo trộn đáp án: "
        "cần lưu và đối chiếu display key qua option_map thay vì chỉ dựa vào thứ tự gốc. "
        "Đã xử lý qua nhiều vòng fix và bổ sung test. Một số migration trùng số (018) "
        "cần theo dõi khi deploy."
    )
    add_h(doc, "6.1. Định hướng công việc tuần sau (gợi ý)", 2)
    add_bullet(doc, "Kiểm thử end-to-end trên môi trường production (GV → SV → chấm → export).")
    add_bullet(doc, "Hoàn thiện báo cáo đồ án Word/PDF; cập nhật sơ đồ và ảnh minh họa từ thesis_assets.")
    add_bullet(doc, "Chuẩn hóa CI chạy test backend/frontend trước khi deploy.")
    add_bullet(doc, "Rà soát bảo mật env, rate limit API công khai.")

    add_h(doc, "7. Kết luận", 1)
    add_p(
        doc,
        "Tuần qua dự án đạt tiến độ rõ rệt: từ bổ sung module quản lý và ngân hàng câu hỏi "
        "đến vận hành được luồng thi online đa mã đề với chấm điểm và giám sát cơ bản. "
        "Hệ thống đã có thể demo đầy đủ vòng đời thi cho báo cáo và bảo vệ, với điều kiện "
        "tiếp tục kiểm thử tích hợp và hoàn thiện tài liệu."
    )

    # Footer margin note
    doc.add_paragraph()
    add_center(
        doc,
        "— Hết báo cáo tuần —",
        bold=False,
        size=12,
    )

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    return doc


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "BaoCaoTuan_11-18_05_2026_OnlineExamination.docx"
    doc = build_report()
    doc.save(out)
    print(f"Đã tạo: {out}")


if __name__ == "__main__":
    main()
