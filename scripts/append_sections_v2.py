"""
Append new Teacher sections (4.4.7–4.4.9) and Admin sections (4.5.7–4.5.8)
into the existing BaoCao_Chuong4_GiaoDien_v2.docx.

- Inserts teacher additions BEFORE the 4.5 admin heading
- Renumbers admin figure captions (shift +4)
- Appends admin additions at the end
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy
import re

INPUT  = r'c:\VS-Code\GraduationProject\docs\BaoCao_Chuong4_GiaoDien_v2.docx'
OUTPUT = r'c:\VS-Code\GraduationProject\docs\BaoCao_Chuong4_GiaoDien_v2.docx'

doc = Document(INPUT)

# ── helpers ──────────────────────────────────────────────────

def make_heading2(doc, text):
    p = doc.add_paragraph(text, style='Heading 2')
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    return p

def make_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p

def make_figure_placeholder(doc, fig_num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f' Hình 4.{fig_num}. {caption}')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p

# ── Step 1: Find insertion point (before 4.5 heading) ───────

admin_heading_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('4.5. Triển khai'):
        admin_heading_idx = i
        break

if admin_heading_idx is None:
    raise ValueError("Could not find 4.5 heading")

print(f"Found 4.5 heading at paragraph index {admin_heading_idx}")

# ── Step 2: Build teacher addition paragraphs ────────────────
# We'll create them via doc.add_paragraph, then move their XML elements

teacher_paras = []

# 4.4.7
teacher_paras.append(('heading', '4.4.7. Giao diện quản lý danh sách sinh viên'))
teacher_paras.append(('body',
    'Giảng viên có thể quản lý danh sách sinh viên thuộc quyền phụ trách thông qua giao diện quản lý sinh viên. '
    'Bảng danh sách hiển thị thông tin: họ tên, username, email, lớp hành chính và trạng thái tài khoản. '
    'Giảng viên có thể tìm kiếm theo từ khóa, chỉnh sửa thông tin cá nhân của sinh viên, '
    'xóa tài khoản và xem mật khẩu tạm thời của sinh viên mới được tạo. '
    'Chức năng gửi email bảng điểm hàng loạt cho phép giảng viên thông báo kết quả '
    'cho toàn bộ hoặc một nhóm sinh viên được chọn.'
))
teacher_paras.append(('figure', 16, 'Giao diện quản lý danh sách sinh viên'))

# 4.4.8
teacher_paras.append(('heading', '4.4.8. Giao diện bảng điểm và transcript sinh viên'))
teacher_paras.append(('body',
    'Tab bảng điểm trong giao diện quản lý sinh viên cho phép giảng viên chọn bài thi, '
    'tìm kiếm sinh viên và xem bảng điểm dạng bảng với thông tin: họ tên, mã đề, điểm số, '
    'trạng thái chấm bài và thời gian nộp bài. '
    'Giảng viên có thể xuất bảng điểm ra file CSV và gửi email thông báo điểm '
    'cho toàn bộ hoặc nhóm sinh viên được chọn.'
))
teacher_paras.append(('figure', 17, 'Giao diện bảng điểm theo bài thi'))

teacher_paras.append(('body',
    'Giảng viên cũng có thể xem bảng điểm tổng hợp (transcript) của từng sinh viên '
    'thông qua modal hiển thị thông tin cá nhân, danh sách các bài thi đã hoàn thành '
    'với điểm số và xếp loại, cùng phần tổng hợp GPA theo thang điểm 4 và thang điểm 10. '
    'Giao diện hỗ trợ in bảng điểm trực tiếp, xuất file HTML hoặc CSV, '
    'và gửi bảng điểm qua email cho sinh viên.'
))
teacher_paras.append(('figure', 18, 'Giao diện bảng điểm tổng hợp (transcript) sinh viên'))

# 4.4.9
teacher_paras.append(('heading', '4.4.9. Giao diện trang chủ giảng viên'))
teacher_paras.append(('body',
    'Trang chủ giảng viên hiển thị bảng tổng quan hoạt động giảng dạy và quản lý thi cử. '
    'Phần thống kê nhanh bao gồm các thẻ số liệu: tổng số sinh viên quản lý, '
    'số bài thi đã tạo, số phiên thi và số bài chờ chấm điểm. '
    'Bên dưới là bảng danh sách sinh viên với chức năng tìm kiếm và phân trang, '
    'cùng phần hoạt động gần đây hiển thị các sự kiện mới nhất như bài thi được nộp, '
    'phiên thi được mở và kết quả chấm bài. '
    'Giảng viên có thể lọc hoạt động theo trạng thái, khoảng thời gian và từ khóa.'
))
teacher_paras.append(('figure', 19, 'Giao diện trang chủ giảng viên'))

# ── Step 3: Create the paragraph XML elements ────────────────

admin_heading_element = doc.paragraphs[admin_heading_idx]._element

new_elements = []
for item in teacher_paras:
    if item[0] == 'heading':
        p = make_heading2(doc, item[1])
    elif item[0] == 'body':
        p = make_body(doc, item[1])
    elif item[0] == 'figure':
        p = make_figure_placeholder(doc, item[1], item[2])
    new_elements.append(p._element)

# Move each new element: remove from end of body, insert before admin heading
body = doc.element.body
for el in new_elements:
    body.remove(el)
    admin_heading_element.addprevious(el)

print(f"Inserted {len(new_elements)} paragraphs before 4.5 heading")

# ── Step 4: Renumber admin figure captions (shift +4) ────────

SHIFT = 4
renumber_count = 0
for p in doc.paragraphs:
    text = p.text.strip()
    m = re.match(r'Hình 4\.(\d+)\.\s', text)
    if m:
        old_num = int(m.group(1))
        if old_num >= 16 and old_num <= 21:
            new_num = old_num + SHIFT
            for run in p.runs:
                if f'4.{old_num}.' in run.text:
                    run.text = run.text.replace(f'4.{old_num}.', f'4.{new_num}.')
                    renumber_count += 1
                    print(f"  Renumbered Hình 4.{old_num} → 4.{new_num}: {run.text.strip()[:60]}")

print(f"Renumbered {renumber_count} figure captions")

# ── Step 5: Append admin additions at the end ────────────────

# 4.5.7
make_heading2(doc, '4.5.7. Giao diện báo cáo hệ thống')
make_body(doc,
    'Giao diện báo cáo hệ thống cung cấp cái nhìn tổng quan về toàn bộ hoạt động của hệ thống '
    'thông qua các chỉ số KPI quan trọng. '
    'Phần thống kê tài khoản hiển thị tổng số tài khoản, số sinh viên, giảng viên và quản trị viên. '
    'Phần thống kê thi cử bao gồm: tổng số bài thi, tổng phiên thi, tỷ lệ hoàn thành, '
    'tỷ lệ đạt và điểm trung bình toàn hệ thống với các thanh tiến trình trực quan.'
)
make_body(doc,
    'Phần thống kê toàn vẹn bài thi hiển thị số vi phạm trong 24 giờ qua, '
    'số phiên bị đánh dấu vi phạm và số bài đang chờ chấm điểm thủ công. '
    'Bên dưới là bảng danh sách các bài thi gần đây với thông tin chi tiết về số phiên, '
    'trạng thái và kết quả tổng hợp, giúp quản trị viên nhanh chóng đánh giá '
    'hiệu quả vận hành tổng thể của hệ thống.'
)
make_figure_placeholder(doc, 26, 'Giao diện báo cáo hệ thống')

# 4.5.8
make_heading2(doc, '4.5.8. Giao diện trang chủ quản trị viên')
make_body(doc,
    'Trang chủ quản trị viên hiển thị bảng tổng quan toàn hệ thống với các thẻ số liệu: '
    'tổng số tài khoản, tổng số bài thi, tổng phiên thi và số lớp hành chính. '
    'Bảng danh sách sinh viên cho phép tìm kiếm và lọc theo lớp hành chính, '
    'giúp quản trị viên nhanh chóng tra cứu thông tin sinh viên.'
)
make_body(doc,
    'Phần hoạt động gần đây liệt kê các sự kiện mới nhất trong hệ thống '
    'như đăng nhập, tạo tài khoản, nộp bài thi và chấm điểm. '
    'Quản trị viên có thể lọc hoạt động theo trạng thái, khoảng thời gian và từ khóa. '
    'Từ trang chủ, quản trị viên có thể nhanh chóng chuyển đến các trang quản lý chi tiết.'
)
make_figure_placeholder(doc, 27, 'Giao diện trang chủ quản trị viên')

# ── Step 6: Save ─────────────────────────────────────────────

doc.save(OUTPUT)
print(f"\nSaved to {OUTPUT}")

# Verify final structure
print("\n=== Final structure ===")
doc2 = Document(OUTPUT)
for p in doc2.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style else ''
    if 'Heading' in style or 'Hình' in text:
        print(f"  {style:12s} | {text[:80]}")
