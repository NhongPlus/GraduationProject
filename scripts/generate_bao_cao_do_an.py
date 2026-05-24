# -*- coding: utf-8 -*-
"""
Sinh báo cáo đồ án Word (~15–25 trang) cho dự án GraduationProject
(Hệ thống thi trực tuyến), bám cấu trúc mẫu đồ án chuẩn (bìa, cam đoan,
cảm ơn, mở đầu, danh mục, mục lục, 5 chương, kết luận, tài liệu tham khảo).

Chạy: python scripts/generate_bao_cao_do_an.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor

from thesis_assets.diagrams import ensure_diagrams


DE_TAI = (
    "XÂY DỰNG HỆ THỐNG THI TRỰC TUYẾN\n"
    "(ONLINE EXAMINATION SYSTEM)"
)


def set_body_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(13)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)


def fmt_para(p, justify: bool = True) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)


def add_center(doc: Document, text: str, *, bold: bool = True, size: int = 13) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    fmt_para(p, justify=False)


def add_p(doc: Document, text: str, *, justify: bool = True) -> None:
    p = doc.add_paragraph(text)
    fmt_para(p, justify=justify)


def add_h(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h.paragraph_format.line_spacing = 1.5


def add_h1_friend(doc: Document, text: str) -> None:
    """Heading 1 giống mẫu Đại Nam (LỜI CAM ĐOAN, LỜI CẢM ƠN, …)."""
    add_h(doc, text, 1)


def add_figure(doc: Document, path: Path, caption: str, *, width_cm: float = 15.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.italic = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    cap.paragraph_format.line_spacing = 1.5
    cap.paragraph_format.space_after = Pt(12)


def add_left_cover_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt_para(p, justify=False)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        fmt_para(p, justify=False)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None) -> None:
    if caption:
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.italic = True
        fmt_para(c, justify=False)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def cover_dai_nam_outer(doc: Document) -> None:
    """Trang bìa ngoài — bám sát mẫu Trường Đại học Đại Nam."""
    add_center(doc, "BỘ GIÁO DỤC ĐÀO TẠO", bold=True, size=14)
    add_center(doc, "TRƯỜNG ĐẠI HỌC ĐẠI NAM", bold=True, size=14)
    doc.add_paragraph()
    add_center(doc, "ĐỒ ÁN TỐT NGHIỆP", bold=True, size=16)
    doc.add_paragraph()
    add_center(doc, DE_TAI, bold=True, size=14)
    for _ in range(3):
        doc.add_paragraph()
    add_left_cover_block(
        doc,
        [
            "SINH VIÊN THỰC HIỆN\t: ................................................",
            "MÃ SINH VIÊN\t: ................................................",
            "KHOA\t: CÔNG NGHỆ THÔNG TIN",
        ],
    )
    doc.add_paragraph()
    add_center(doc, "HÀ NỘI 2026", bold=True, size=14)


def cover_dai_nam_inner(doc: Document) -> None:
    """Trang bìa lặp trong (như mẫu bạn — khối tên + đề tài + Hà Nội)."""
    add_center(doc, "BỘ GIÁO DỤC ĐÀO TẠO", bold=True, size=14)
    add_center(doc, "TRƯỜNG ĐẠI HỌC ĐẠI NAM", bold=True, size=14)
    add_center(doc, "------------------------------", bold=False, size=12)
    add_center(doc, "................................", bold=True, size=14)
    add_center(doc, DE_TAI, bold=True, size=13)
    for _ in range(4):
        doc.add_paragraph()
    add_center(doc, "HÀ NỘI 2026", bold=True, size=14)


def front_matter(doc: Document) -> None:
    cover_dai_nam_outer(doc)
    doc.add_page_break()
    cover_dai_nam_inner(doc)
    doc.add_page_break()

    add_h1_friend(doc, "LỜI CAM ĐOAN")
    add_p(
        doc,
        "Tôi xin cam đoan rằng đồ án tốt nghiệp với đề tài “Xây dựng hệ thống thi trực tuyến "
        "(Online Examination System)” là kết quả nghiên cứu và thực hiện của cá nhân tôi dưới sự "
        "hướng dẫn của giảng viên hướng dẫn.",
    )
    add_p(
        doc,
        "Các nội dung phân tích, thiết kế, triển khai và đánh giá được trình bày trong đồ án là "
        "trung thực; phần tham khảo từ tài liệu, mã nguồn mở và các công trình khác đã được trích "
        "dẫn trong phần tài liệu tham khảo.",
    )
    add_p(
        doc,
        "Tôi xin hoàn toàn chịu trách nhiệm trước nhà trường về tính trung thực của nội dung đồ án.",
    )
    p_sign = doc.add_paragraph("Hà Nội, ngày ___ tháng ___ năm 2026")
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fmt_para(p_sign, justify=False)
    add_p(doc, "\nSinh viên\n(ký và ghi rõ họ tên)", justify=False)
    doc.add_page_break()

    add_h1_friend(doc, "LỜI CẢM ƠN")
    add_p(
        doc,
        "Trong quá trình thực hiện đồ án, tôi đã nhận được sự hỗ trợ của giảng viên hướng dẫn, "
        "quý thầy cô trong khoa và gia đình, bạn bè.",
    )
    add_p(
        doc,
        "Tôi xin bày tỏ lòng biết ơn tới giảng viên hướng dẫn đã định hướng đề tài, góp ý kiến "
        "trọng yếu giúp hệ thống hoàn thiện hơn về kiến trúc, kiểm thử và trình bày báo cáo.",
    )
    add_p(
        doc,
        "Mặc dù đã cố gắng hoàn thành đồ án, đề tài vẫn có thể còn hạn chế. Tôi rất mong nhận "
        "thêm ý kiến đóng góp để cải thiện trong các phiên bản tiếp theo.",
    )
    add_p(doc, "Xin chân thành cảm ơn.")
    doc.add_page_break()

    add_h1_friend(doc, "LỜI NÓI ĐẦU")
    add_p(
        doc,
        "Trong bối cảnh chuyển đổi số giáo dục, hình thức kiểm tra – đánh giá trực tuyến ngày càng "
        "phổ biến. Đại dịch COVID-19 đã thúc đẩy nhanh việc tổ chức thi trên môi trường mạng, đặt "
        "ra yêu cầu về tính ổn định, bảo mật, công bằng và khả năng giám sát hành vi thí sinh.",
    )
    add_p(
        doc,
        "Một hệ thống thi trực tuyến hoàn chỉnh không chỉ là trang web hiển thị câu hỏi, mà cần "
        "chuỗi nghiệp vụ: quản trị người dùng và phân quyền, quản lý môn học và đề thi, tổ chức "
        "phiên thi, đồng bộ thời gian, tự động lưu bài, chống gian lận ở mức hợp lý, kênh thông "
        "báo thời gian thực cho giám thị, chấm điểm và thống kê kết quả.",
    )
    add_p(
        doc,
        "Đề tài “Xây dựng hệ thống thi trực tuyến” được thực hiện trên mã nguồn dự án "
        "GraduationProject: frontend React/TypeScript, backend Node.js/Express, cơ sở dữ liệu "
        "PostgreSQL, xác thực JWT, Socket.IO cho tín hiệu thi, cùng các tài liệu OpenAPI và hướng "
        "dẫn triển khai trong kho mã nguồn.",
    )
    add_p(
        doc,
        "Báo cáo được tổ chức thành năm chương: tổng quan; cơ sở lý thuyết và công nghệ; phân tích "
        "thiết kế; triển khai kiểm thử; đánh giá và hướng phát triển, kèm kết luận và tài liệu tham khảo.",
    )
    doc.add_page_break()

    add_h1_friend(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
    add_p(doc, "………………………………………………………………………………………………………………", justify=False)
    add_p(doc, "………………………………………………………………………………………………………………", justify=False)
    add_p(doc, "………………………………………………………………………………………………………………", justify=False)
    doc.add_page_break()


def danh_muc(doc: Document) -> None:
    add_center(doc, "DANH MỤC VIẾT TẮT", bold=True, size=14)
    add_table(
        doc,
        ["Viết tắt", "Ý nghĩa"],
        [
            ["API", "Application Programming Interface"],
            ["BE", "Backend"],
            ["CEFR", "Khung tham chiếu (tham khảo trong bối cảnh giáo dục — không dùng trực tiếp trong đề tài)"],
            ["CRUD", "Create, Read, Update, Delete"],
            ["FE", "Frontend"],
            ["HTTP/HTTPS", "Giao thức truyền tải siêu văn bản"],
            ["JWT", "JSON Web Token"],
            ["MCQ", "Multiple Choice Question (câu hỏi trắc nghiệm)"],
            ["ORM", "Object–Relational Mapping (khái niệm tổng quát; dự án dùng truy vấn SQL trực tiếp)"],
            ["RBAC", "Role-Based Access Control"],
            ["REST", "Representational State Transfer"],
            ["SMTP", "Simple Mail Transfer Protocol"],
            ["SPA", "Single Page Application"],
            ["SQL", "Structured Query Language"],
            ["UI/UX", "Giao diện / trải nghiệm người dùng"],
            ["UUID", "Định danh duy nhất toàn cục"],
        ],
        None,
    )
    doc.add_page_break()

    add_center(doc, "DANH MỤC HÌNH ẢNH", bold=True, size=14)
    hinh = [
        "Hình 2.1. Kiến trúc tổng thể: trình duyệt — API — PostgreSQL và Socket.IO (đã chèn trong báo cáo)",
        "Hình 2.2. Mô hình REST — HTTP và JSON (đã chèn)",
        "Hình 2.3. Luồng JWT và RBAC (đã chèn)",
        "Hình 2.4. Socket.IO — room và sự kiện (đã chèn)",
        "Hình 2.5. Thực thể dữ liệu khái niệm (đã chèn)",
        "Hình 2.6. Mô hình React — UI / state / side-effects (đã chèn)",
        "Hình 2.7. Vòng đời phiên thi (đã chèn)",
        "Hình 3.1. Sơ đồ use case tổng thể (đã chèn)",
        "Hình 3.2. Sơ đồ ER quan hệ thực thể (đã chèn)",
        "Hình 4.x — Ảnh cấu trúc thư mục, Swagger, màn hình triển khai (tự chèn)",
    ]
    for h in hinh:
        add_p(doc, h, justify=False)
    doc.add_page_break()

    add_center(doc, "DANH MỤC BẢNG BIỂU", bold=True, size=14)
    bang = [
        "Bảng 2.1. So sánh nền tảng thi trực tuyến phổ biến",
        "Bảng 2.2. Bảng công nghệ theo tầng hệ thống",
        "Bảng 3.1. Danh sách tác nhân",
        "Bảng 3.2. Ma trận yêu cầu chức năng theo vai trò",
        "Bảng 3.3. Đặc tả use case — Sinh viên: đăng nhập, xem đề, làm bài",
        "Bảng 3.4. Đặc tả use case — Giáo viên: tạo phiên, giám sát, chấm",
        "Bảng 3.5. Đặc tả use case — Quản trị: người dùng, audit, cấu hình",
        "Bảng 3.6. Ma trận khóa ngoại (logic)",
        "Bảng 3.7. Thuộc tính chính theo bảng",
        "Bảng 3.8. Tóm tắt thực thể — PK và hướng liên kết",
        "Bảng 4.1. Môi trường phát triển và phiên bản công cụ",
        "Bảng 4.2. Danh mục kiểm thử thủ công theo kịch bản",
    ]
    for b in bang:
        add_p(doc, b, justify=False)
    doc.add_page_break()


def muc_luc(doc: Document) -> None:
    add_center(doc, "MỤC LỤC", bold=True, size=14)
    lines = [
        "CHƯƠNG 1: TỔNG QUAN VỀ ĐỀ TÀI",
        "    1.1. Giới thiệu chung và tính cấp thiết của đề tài",
        "    1.2. Mục tiêu nghiên cứu",
        "    1.3. Phạm vi và đối tượng nghiên cứu",
        "    1.4. Phương pháp nghiên cứu",
        "    1.5. Cấu trúc đồ án",
        "    1.6. Ý nghĩa thực tiễn và học thuật",
        "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG",
        "    2.1. Tổng quan e-assessment và kiến trúc phần mềm",
        "    2.2. Chuẩn hóa REST và mô hình tài nguyên",
        "    2.3. Nghiên cứu sản phẩm tương tự và định vị đồ án",
        "    2.4. JWT, RBAC và an toàn phiên làm bài",
        "    2.5. Socket.IO trong phòng thi số",
        "    2.6. PostgreSQL, migration và dữ liệu quan hệ",
        "    2.7. React, TypeScript, Vite và UX làm bài",
        "    2.8. Vòng đời phiên thi và import Word",
        "    2.9. i18n và khả năng truy cập (a11y)",
        "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
        "    3.1. Phân tích yêu cầu chức năng và phi chức năng",
        "    3.2. Các tác nhân và ma trận chức năng",
        "    3.3. Phân tích use case theo vai trò",
        "    3.4. Thiết kế luồng xử lý chính",
        "    3.5. Thiết kế API REST và tài liệu OpenAPI",
        "    3.6. Thiết kế dữ liệu: bảng quan hệ, ma trận FK và sơ đồ ER",
        "    3.7. Autosave, integrity và phục hồi runtime",
        "    3.8. Thiết kế giám sát thi và sự kiện Socket",
        "CHƯƠNG 4: TRIỂN KHAI VÀ KIỂM THỬ HỆ THỐNG",
        "    4.1. Môi trường phát triển và công cụ",
        "    4.2. Cấu trúc thư mục dự án",
        "    4.3. Triển khai chức năng theo vai trò",
        "    4.4. Kiểm thử đơn vị và kiểm thử hệ thống",
        "    4.5. Triển khai môi trường thử nghiệm và production",
        "    4.6. Hiệu năng, log và xử lý sự cố",
        "NỘI DUNG BỔ SUNG: MÔ TẢ CHI TIẾT LUỒNG NGHIỆP VỤ, VẬN HÀNH VÀ BẢO MẬT",
        "CHƯƠNG 5: ĐÁNH GIÁ KẾT QUẢ VÀ HƯỚNG PHÁT TRIỂN",
        "    5.1. Kết quả đạt được",
        "    5.2. Hạn chế",
        "    5.3. Hướng phát triển",
        "PHỤ LỤC: DANH MỤC KIỂM THỬ CHI TIẾT (GỢI Ý)",
        "KẾT LUẬN",
        "TÀI LIỆU THAM KHẢO",
    ]
    for ln in lines:
        add_p(doc, ln, justify=False)
    doc.add_page_break()


def chuong1(doc: Document) -> None:
    add_h(doc, "CHƯƠNG 1: TỔNG QUAN VỀ ĐỀ TÀI", 1)
    add_h(doc, "1.1. Giới thiệu chung và tính cấp thiết của đề tài", 2)
    add_p(
        doc,
        "Kiểm tra và đánh giá kết quả học tập là khâu then chốt trong quản lý đào tạo. Khi quy mô "
        "lớp học tăng và học viên phân tán địa lý, tổ chức thi tập trung tại phòng máy gặp hạn chế "
        "về chi phí cơ sở vật chất, lịch thi và khả năng mở rộng. Thi trực tuyến cho phép linh hoạt "
        "thời gian, giảm chi phí in ấn đề giấy, tự động hóa một phần chấm trắc nghiệm và thống kê "
        "nhanh kết quả.",
    )
    add_p(
        doc,
        "Tuy nhiên, thi trực tuyến đặt ra thách thức về toàn vẹn dữ liệu bài làm, đồng bộ thời gian "
        "làm bài, phát hiện hành vi bất thường, và phân quyền chặt chẽ để tránh lộ đề hoặc truy cập "
        "trái phép. Do đó, việc xây dựng một hệ thống có kiến trúc rõ ràng, có tài liệu API và quy "
        "ước autosave/integrity là cần thiết cho cả mục đích học tập và làm cơ sở mở rộng thương mại.",
    )
    add_p(
        doc,
        "Xuất phát từ nhu cầu trên, đồ án lựa chọn đề tài xây dựng hệ thống thi trực tuyến với đầy "
        "đủ vai trò quản trị viên, giáo viên và sinh viên, tích hợp import đề Word, timer đồng bộ "
        "máy chủ, autosave, Socket.IO cho giám sát, và các chức năng thống kê — phù hợp với xu "
        "hướng ứng dụng web hiện đại.",
    )

    add_h(doc, "1.2. Mục tiêu nghiên cứu", 2)
    add_p(
        doc,
        "Mục tiêu tổng quát: xây dựng phần mềm web hỗ trợ vòng đời tổ chức thi trực tuyến, có khả "
        "năng triển khai cục bộ phục vụ demo và có thể cấu hình triển khai môi trường thật.",
    )
    add_p(doc, "Mục tiêu cụ thể:")
    add_bullets(
        doc,
        [
            "Phân tích yêu cầu và thiết kế use case theo từng vai trò.",
            "Triển khai backend REST /v1 với PostgreSQL, migration schema, JWT và RBAC.",
            "Triển khai frontend SPA với React, quản lý trạng thái, định tuyến và i18n.",
            "Hiện thực làm bài thi: timer, autosave, integrity events, fullscreen theo cấu hình đề.",
            "Hiện thực Socket.IO cho tín hiệu thi (bắt đầu, cảnh báo, force-submit) và màn hình giám thị.",
            "Bổ sung tài liệu OpenAPI và hướng dẫn kiểm thử nhanh trong kho mã nguồn.",
        ],
    )

    add_h(doc, "1.3. Phạm vi và đối tượng nghiên cứu", 2)
    add_p(
        doc,
        "Phạm vi đồ án tập trung vào hệ thống trong kho GraduationProject: module người dùng, đề "
        "thi, phiên thi, làm bài, chấm điểm, giám sát, thống kê và các tiện ích liên quan được mô "
        "tả trong DO_AN_MASTER.md. Phạm vi không bao gồm phần cứng phòng thi, sinh trắc học nâng "
        "cao, hay chứng thực pháp lý điện tử đầy đủ cấp quốc gia.",
    )
    add_p(
        doc,
        "Đối tượng nghiên cứu là kiến trúc ứng dụng web đa vai trò, các mẫu thiết kế API, quản lý "
        "phiên làm bài an toàn, và trải nghiệm người dùng khi làm bài dài với mạng không ổn định.",
    )

    add_h(doc, "1.4. Phương pháp nghiên cứu", 2)
    add_bullets(
        doc,
        [
            "Khảo sát tài liệu: đọc mã nguồn, API.md, openapi.yaml, tài liệu contract autosave/integrity.",
            "Phân tích so sánh: đối chiếu với các nền tảng phổ biến (Moodle, Google Forms, hệ thống thương mại).",
            "Thực nghiệm: cài đặt cục bộ, chạy migration, thực hiện kịch bản kiểm thử thủ công end-to-end.",
            "Đánh giá: kiểm tra log, hành vi Socket, và test tự động phần backend (Vitest).",
        ],
    )

    add_h(doc, "1.5. Cấu trúc đồ án", 2)
    add_p(
        doc,
        "Đồ án gồm năm chương: (1) tổng quan; (2) cơ sở lý thuyết và công nghệ; (3) phân tích và "
        "thiết kế; (4) triển khai và kiểm thử; (5) đánh giá và hướng phát triển; kết luận và tài "
        "liệu tham khảo.",
    )
    add_h(doc, "1.6. Ý nghĩa thực tiễn và học thuật", 2)
    add_p(
        doc,
        "Về thực tiễn, hệ thống có thể phục vụ các khóa học ngắn hạn, kiểm tra giữa kỳ trực tuyến "
        "hoặc thi thử trong trường đại học, giúp giảm khối lượng vận hành giấy tờ và rút ngắn thời "
        "gian công bố điểm cho phần trắc nghiệm. Khi kết hợp quy trình giám sát con người và camera "
        "ngoài phạm vi đồ án, có thể tiến tới mô hình phòng thi lai (hybrid).",
    )
    add_p(
        doc,
        "Về học thuật, đồ án cho phép sinh viên ôn luyện các môn CSDL, Lập trình web, Phần mềm mã "
        "nguồn mở, An toàn thông tin cơ bản và Quản trị mạng máy tính thông qua một case study thống "
        "nhất. Việc đọc OpenAPI và viết kịch bản kiểm thử cũng rèn kỹ năng tư duy hệ thống.",
    )
    add_p(
        doc,
        "Đồng thời, đề tài nhấn mạnh kỹ năng viết tài liệu kỹ thuật: mô tả use case, ma trận quyền, "
        "hướng dẫn triển khai và phân tích rủi ro — những nội dung thường xuất hiện trong quy trình "
        "phát triển phần mềm chuyên nghiệp.",
    )


def chuong2(doc: Document, img: dict[str, Path]) -> None:
    add_h(doc, "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 1)

    add_h(doc, "2.1. Tổng quan đánh giá trực tuyến (e-assessment) và kiến trúc phần mềm", 2)
    add_p(
        doc,
        "Trong các hệ thống đào tạo hiện đại, đánh giá năng lực người học không chỉ dừng ở bài kiểm "
        "tra trên giấy mà còn được thực hiện trên môi trường số. Đánh giá trực tuyến (e-assessment) "
        "là quá trình tổ chức đề thi, giám sát phiên làm bài, thu thập bài làm, lưu trữ dữ liệu và "
        "chấm điểm thông qua các dịch vụ phần mềm. Mô hình phổ biến là ứng dụng web nhiều lớp: "
        "trình duyệt đóng vai trò client, máy chủ ứng dụng triển khai nghiệp vụ và API, cơ sở dữ "
        "liệu quan hệ đảm bảo tính bền vững của dữ liệu.",
    )
    add_p(
        doc,
        "Đồ án “Hệ thống thi trực tuyến” đi theo đúng mô hình đó: React/TypeScript cho lớp giao diện, "
        "Node.js/Express cho lớp dịch vụ REST, PostgreSQL cho lớp lưu trữ. Bên cạnh kênh HTTP "
        "truyền thống, hệ thống bổ sung Socket.IO để phục vụ giám sát thời gian thực và phát tín "
        "hiệu bắt đầu/kết thúc thi — đây là điểm khác biệt so với các form khảo sát đơn giản.",
    )
    add_p(
        doc,
        "Một thách thức mang tính hệ thống là phải đồng thời đảm bảo trải nghiệm người dùng mượt "
        "(autosave, đồng hồ đếm ngược) và tính an toàn (phân quyền, chống đọc trộm đáp án, ghi log "
        "hành vi). Hình 2.1 minh họa phân tách lớp và luồng chính giữa các thành phần.",
    )
    add_figure(
        doc,
        img["h2_01"],
        "Hình 2.1. Kiến trúc tổng thể: trình duyệt — API — PostgreSQL và kênh Socket.IO",
    )

    add_h(doc, "2.2. Chuẩn hóa giao tiếp REST và mô hình tài nguyên", 2)
    add_p(
        doc,
        "REST (Representational State Transfer) là phong cách thiết kế API sử dụng các phương thức "
        "HTTP một cách nhất quán: GET để đọc, POST để tạo, PATCH/PUT để cập nhật, DELETE để xóa. "
        "Tài nguyên được đặt tên bằng danh từ (ví dụ /v1/exams) và biểu diễn bằng JSON. Mã trạng "
        "thái HTTP (200, 201, 400, 401, 403, 404, 409, 500) giúp client phân loại lỗi và hiển thị "
        "thông báo phù hợp cho người dùng cuối.",
    )
    add_p(
        doc,
        "Trong đồ án, tiền tố phiên bản /v1 giúp tách contract: khi có thay đổi lớn có thể triển khai "
        "/v2 song song mà không phá vỡ client cũ. Việc duy trì file openapi.yaml giúp đội FE/BE "
        "thống nhất schema request/response, giảm lỗi tích hợp khi số lượng endpoint tăng.",
    )
    add_figure(
        doc,
        img["h2_02"],
        "Hình 2.2. Mô hình REST — tài nguyên, phương thức HTTP và biểu diễn JSON",
    )

    add_h(doc, "2.3. Nghiên cứu sản phẩm tương tự và định vị đồ án", 2)
    add_table(
        doc,
        ["Nền tảng", "Ưu điểm", "Nhược điểm / ghi chú"],
        [
            ["Moodle Quiz", "Phổ biến trong giáo dục, plugin phong phú", "Cài đặt nặng, tùy biến UI phức tạp"],
            ["Google Forms", "Nhanh, dễ dùng", "Hạn chế proctoring, ít kiểm soát server-side timer nâng cao"],
            ["Hệ thống thương mại", "Proctoring mạnh, SLA", "Chi phí, khó tùy chỉnh mã nguồn"],
            ["Đồ án (GraduationProject)", "Mã nguồn mở, stack hiện đại, OpenAPI", "Cần hardening production, CI đầy đủ"],
        ],
        "Bảng 2.1. So sánh nhanh các hướng tiếp cận",
    )
    add_p(
        doc,
        "Bảng trên cho thấy đồ án hướng tới giải pháp “cân bằng” giữa tính năng chuyên sâu của LMS "
        "và tốc độ triển khai của ứng dụng web tùy chỉnh. Điểm mạnh của đồ án là minh bạch mã nguồn, "
        "dễ demo cục bộ và mở rộng theo nhánh đề tài (AI dự đoán điểm, export, ngân hàng câu hỏi…).",
    )

    add_h(doc, "2.4. Xác thực JWT, RBAC và an toàn phiên làm bài", 2)
    add_p(
        doc,
        "JWT cho phép máy chủ xác thực request mà không cần lưu session server-side cổ điển. Token "
        "chứa các claim tối thiểu (định danh, vai trò, thời hạn) và được ký bằng HMAC hoặc RSA. "
        "RBAC (Role-Based Access Control) gắn tập quyền theo admin/teacher/student, được thực thi "
        "ở lớp middleware sau khi token đã hợp lệ.",
    )
    add_p(
        doc,
        "Với thi trực tuyến, lớp “toàn vẹn phiên thi” còn bao gồm autosave định kỳ, lưu runtime để "
        "khôi phục timer khi máy chủ restart, và integrity events (ví dụ thoát fullscreen). Các cơ "
        "chế này không thay thế giám sát con người nhưng tạo dữ liệu khách quan phục vụ quyết định "
        "của giám thị.",
    )
    add_figure(
        doc,
        img["h2_03"],
        "Hình 2.3. Luồng xác thực JWT và kiểm tra quyền (RBAC) sau xác thực",
    )

    add_h(doc, "2.5. Giao tiếp thời gian thực với Socket.IO trong phòng thi số", 2)
    add_p(
        doc,
        "Socket.IO cung cấp kênh song song với HTTP, phù hợp đẩy sự kiện tức thời: bắt đầu thi, "
        "cảnh báo, buộc nộp bài, cập nhật danh sách thí sinh online. Mô hình room theo examId giúp "
        "cô lập sự kiện giữa các ca thi. Khi triển khai nhiều instance, cần Redis adapter để đồng "
        "bộ trạng thái socket giữa các node.",
    )
    add_figure(
        doc,
        img["h2_04"],
        "Hình 2.4. Mô hình Socket.IO — giám thị, server hub và nhóm thí sinh trong room examId",
    )

    add_h(doc, "2.6. PostgreSQL, migration và thiết kế dữ liệu quan hệ", 2)
    add_p(
        doc,
        "PostgreSQL hỗ trợ transaction ACID, khóa ngoại, index và các kiểu dữ liệu nâng cao (JSONB). "
        "Với nghiệp vụ thi, transaction giúp gom các thao tác “khóa phiên — ghi bài nộp — cập nhật "
        "điểm trắc nghiệm” thành một đơn vị nguyên tố, tránh trạng thái dở khi lỗi giữa chừng.",
    )
    add_p(
        doc,
        "Migration trong kho mã nguồn cho phép tái tạo schema trên môi trường mới. Thực hành tốt là "
        "luôn chạy migration trên bản sao trước khi production, đo thời gian lock bảng lớn và có "
        "kế hoạch index sau khi dữ liệu đã ổn định.",
    )
    add_figure(
        doc,
        img["h2_05"],
        "Hình 2.5. Mô hình thực thể khái niệm: users — exams — questions — sessions — submissions",
    )

    add_h(doc, "2.7. React, TypeScript, Vite và trải nghiệm giao diện làm bài", 2)
    add_table(
        doc,
        ["Tầng", "Công nghệ chính trong GraduationProject"],
        [
            ["Frontend", "React 19, TypeScript, Vite 7, Mantine v8, Redux Toolkit, React Router v7, i18next, Axios"],
            ["Backend", "Node.js, TypeScript, Express 5, Joi, bcrypt, Nodemailer, Mammoth (docx)"],
            ["CSDL", "PostgreSQL (ví dụ Neon trên cloud)"],
            ["Tài liệu API", "openapi.yaml, Swagger /docs (khi bật)"],
        ],
        "Bảng 2.2. Công nghệ theo tầng",
    )
    add_p(
        doc,
        "React chia UI thành component, kết hợp TypeScript để bắt lỗi kiểu sớm. Vite rút ngắn thời "
        "gian khởi động dev server. Mantine cung cấp component đồng nhất; Redux Toolkit giúp "
        "quản lý state toàn cục (phiên đăng nhập, snapshot đề). Với màn hình làm bài, UI phải tối "
        "ưu re-render khi đồng hồ tick và khi autosave phản hồi.",
    )
    add_figure(
        doc,
        img["h2_06"],
        "Hình 2.6. Mô hình React — component UI, state và tác vụ mạng (Axios / Socket client)",
    )

    add_h(doc, "2.8. Vòng đời phiên thi, autosave và import đề Word", 2)
    add_p(
        doc,
        "Vòng đời phiên thi đi từ khâu soạn đề, mở session, sinh viên làm bài với autosave, tới "
        "nộp bài hoặc hết giờ, rồi chấm điểm và thống kê. Hình 2.7 tóm tắt trực quan các bước chính. "
        "Import đề Word dùng Mammoth để chuyển docx → HTML/JSON preview, giúp giáo viên kiểm tra "
        "trước khi commit vào CSDL.",
    )
    add_figure(
        doc,
        img["h2_07"],
        "Hình 2.7. Vòng đời phiên thi (runtime) từ tạo đề tới chấm điểm",
    )
    add_p(
        doc,
        "Rủi ro khi import file lớn cần được kiểm soát bằng giới hạn kích thước và timeout; đồng "
        "thời nên log lỗi parse để hỗ trợ người soạn đề chỉnh lại mẫu Word chuẩn hóa.",
    )

    add_h(doc, "2.9. i18n, UX làm bài thi và khả năng truy cập (a11y)", 2)
    add_p(
        doc,
        "i18next hỗ trợ đa ngôn ngữ (vi/en/ja), giúp hệ thống tiếp cận người dùng đa dạng. UX làm "
        "bài cần đồng hồ rõ ràng, cảnh báo trước khi nộp, thông báo autosave, và xử lý mất mạng "
        "một cách êm ái (retry backoff).",
    )
    add_p(
        doc,
        "Khả năng truy cập nên được cải thiện dần: thông báo thời gian còn lại bằng aria-live, "
        "đảm bảo thao tác bàn phím cho các nút quan trọng, tránh chỉ dùng màu sắc để báo trạng thái "
        "đúng/sai.",
    )


def chuong3(doc: Document, img: dict[str, Path]) -> None:
    add_h(doc, "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    add_h(doc, "3.1. Phân tích yêu cầu chức năng và phi chức năng", 2)
    add_h(doc, "3.1.1. Yêu cầu chức năng", 3)
    add_p(
        doc,
        "Theo vai trò: quản trị viên quản lý người dùng, môn học, reset mật khẩu, audit, báo cáo "
        "hệ thống; giáo viên tổ chức đề thi, phiên thi, giám sát, chấm tự luận; sinh viên đăng nhập, "
        "làm bài, nộp bài, xem kết quả và phần giải thích nếu được phép.",
    )
    add_p(
        doc,
        "Chức năng bổ trợ: import đề Word (preview/commit), ngân hàng câu hỏi, export kết quả, xáo "
        "trộn câu theo chương, dự đoán điểm AI (tùy cấu hình MiniMax).",
    )
    add_h(doc, "3.1.2. Yêu cầu phi chức năng", 3)
    add_bullets(
        doc,
        [
            "Bảo mật: HTTPS production, JWT_SECRET mạnh, không commit secret.",
            "Hiệu năng: tối ưu truy vấn, autosave không block UI.",
            "Khả dụng: giao diện rõ ràng, đa ngôn ngữ (vi/en/ja).",
            "Tin cậy: migration DB đồng bộ schema; log lỗi server.",
            "Mở rộng: tách module router, tài liệu OpenAPI.",
        ],
    )

    add_h(doc, "3.2. Các tác nhân và ma trận chức năng", 2)
    add_table(
        doc,
        ["Tác nhân", "Mô tả"],
        [
            ["Quản trị viên (admin)", "Quản lý người dùng, cấu hình nhạy cảm, audit, giám sát proctoring"],
            ["Giáo viên (teacher)", "Soạn đề/phiên, runtime thi, chấm điểm, thống kê lớp"],
            ["Sinh viên (student)", "Làm bài, autosave, nộp bài, xem kết quả"],
        ],
        "Bảng 3.1. Tóm tắt tác nhân",
    )
    add_table(
        doc,
        ["Chức năng", "Admin", "Teacher", "Student"],
        [
            ["Tạo/sửa đề (một số route)", "Có (theo FE route)", "Theo API", "Không"],
            ["Làm bài / autosave", "Theo cấu hình demo", "Theo cấu hình", "Có"],
            ["Giám sát proctoring", "Có", "Theo API", "Không"],
            ["Chấm điểm tự luận", "Theo API", "Có", "Không"],
        ],
        "Bảng 3.2. Ma trận thể hiện tóm tắt (chi tiết xem ROLES_AND_PERMISSIONS.md và openapi)",
    )
    add_p(
        doc,
        "Hình 3.1 minh họa quan hệ giữa ba tác nhân chính và khối chức năng hệ thống (REST + Socket). "
        "Đây là sơ đồ khối — chi tiết từng use case có thể bổ sung thêm trang sau hoặc phụ lục.",
    )
    add_figure(
        doc,
        img["h3_02"],
        "Hình 3.1. Sơ đồ use case tổng thể theo tác nhân (khối nghiệp vụ)",
        width_cm=15.5,
    )

    add_h(doc, "3.3. Phân tích use case theo vai trò", 2)
    add_table(
        doc,
        ["Mã UC", "Tên", "Mô tả tóm tắt", "Luồng chính"],
        [
            ["SV-01", "Đăng nhập", "Xác thực JWT", "Nhập email/mật khẩu → BE trả token"],
            ["SV-02", "Xem danh sách đề", "Lọc đề được phép", "GET /v1/exams"],
            ["SV-03", "Vào làm bài", "Fullscreen + timer", "Tải đề, restore runtime nếu có"],
            ["SV-04", "Autosave", "Lưu nháp định kỳ", "POST autosave theo contract"],
            ["SV-05", "Gửi integrity", "Ghi nhận vi phạm", "POST integrity-events"],
            ["SV-06", "Nộp bài", "Kết thúc phiên", "POST submit"],
            ["SV-07", "Xem kết quả", "Điểm + đáp án", "GET submission/result theo quyền"],
        ],
        "Bảng 3.3. Đặc tả rút gọn — Sinh viên",
    )
    add_table(
        doc,
        ["Mã UC", "Tên", "Mô tả tóm tắt", "Luồng chính"],
        [
            ["GV-01", "Tạo phiên thi", "Runtime", "POST start-runtime"],
            ["GV-02", "Giám sát", "Socket + log", "Theo dõi integrity, broadcast"],
            ["GV-03", "Buộc nộp", "Kết thúc sớm", "POST force-submit"],
            ["GV-04", "Chấm điểm", "Tự luận", "PATCH grade"],
            ["GV-05", "Import Word", "Soạn đề", "import-word preview/commit"],
        ],
        "Bảng 3.4. Đặc tả rút gọn — Giáo viên",
    )
    add_table(
        doc,
        ["Mã UC", "Tên", "Mô tả tóm tắt", "Luồng chính"],
        [
            ["AD-01", "CRUD user", "Quản trị", "/v1/users"],
            ["AD-02", "Audit log", "Truy vết", "route admin audit"],
            ["AD-03", "Reset password", "Hỗ trợ", "admin password-resets"],
            ["AD-04", "System report", "Vận hành", "admin system-report"],
        ],
        "Bảng 3.5. Đặc tả rút gọn — Quản trị",
    )
    add_p(
        doc,
        "Ghi chú triển khai frontend: một số màn hình chỉ dành cho admin theo route authority; "
        "giáo viên và sinh viên dùng chung nhãn 'user' ở FE nhưng backend vẫn phân biệt role để "
        "áp dụng middleware — chi tiết trong tài liệu ROLES_AND_PERMISSIONS.md.",
    )

    add_h(doc, "3.4. Thiết kế luồng xử lý chính", 2)
    add_p(
        doc,
        "Luồng làm bài: sau khi xác thực, sinh viên mở đề → máy chủ khởi tạo/khôi phục runtime → "
        "đồng hồ đếm ngược hiển thị đồng bộ; trong quá trình làm bài, client gửi autosave và có thể "
        "gửi integrity events; khi hết giờ hoặc nộp bài, máy chủ khóa phiên và tính điểm phần trắc nghiệm.",
    )
    add_p(
        doc,
        "Luồng giám sát: giáo viên/admin mở trang proctoring, lắng nghe Socket, xem danh sách thí "
        "sinh online, gửi cảnh báo hoặc force-submit khi cần quyết định nghiệp vụ.",
    )

    add_h(doc, "3.5. Thiết kế API REST và tài liệu OpenAPI", 2)
    add_p(
        doc,
        "API gốc /v1 nhóm theo router: auth, users, exams, exam-sessions, subjects, question-bank, "
        "v.v. OpenAPI mô tả contract giúp frontend và công cụ sinh client đồng bộ. Các endpoint nhạy "
        "cảm được bảo vệ bởi middleware kiểm tra role.",
    )
    add_p(
        doc,
        "Nguyên tắt thiết kế REST được áp dụng: danh từ số nhiều cho tài nguyên, mã trạng thái HTTP "
        "chuẩn (200/201/400/401/403/404/409/500), thông điệp lỗi thống nhất. Versioning /v1 giúp "
        "mở rộng v2 trong tương lai mà không phá vỡ client cũ nếu duy trì song song.",
    )

    add_h(doc, "3.6. Thiết kế dữ liệu: bảng quan hệ, ma trận khóa ngoại và sơ đồ ER", 2)
    add_p(
        doc,
        "Thiết kế dữ liệu quan hệ thể hiện rõ các thực thể, khóa chính (PK), khóa ngoại (FK) và "
        "bản số (cardinality). Bảng 3.6 liệt kê ma trận FK — thuận tiện khi viết migration và kiểm "
        "tra toàn vẹn tham chiếu. Bảng 3.7 mô tả thuộc tính chính (rút gọn) để độc giả nắm schema "
        "mà không cần mở mã nguồn. Hình 3.2 là sơ đồ ER trực quan có ghi nhãn quan hệ 1–N.",
    )
    add_table(
        doc,
        ["Bảng con", "Cột FK", "Bảng cha", "Cột PK tham chiếu", "Cardinality"],
        [
            ["exams", "author_id", "users", "id", "N : 1"],
            ["exams", "subject_id", "subjects", "id", "N : 1"],
            ["questions", "exam_id", "exams", "id", "N : 1"],
            ["exam_sessions", "exam_id", "exams", "id", "N : 1"],
            ["submissions", "session_id", "exam_sessions", "id", "N : 1"],
            ["submissions", "user_id", "users", "id", "N : 1"],
            ["integrity_events", "session_id", "exam_sessions", "id", "N : 1"],
            ["integrity_events", "user_id", "users", "id", "N : 1"],
        ],
        "Bảng 3.6. Ma trận khóa ngoại (quan hệ logic giữa các bảng)",
    )
    add_table(
        doc,
        ["Bảng", "Thuộc tính", "Kiểu (gợi ý)", "Ghi chú"],
        [
            ["users", "id, email, password_hash, role, created_at", "UUID/serial, text, text, enum, timestamptz", "email unique"],
            ["subjects", "id, name, code", "serial, text, text", "code có thể unique"],
            ["exams", "id, author_id, subject_id, title, duration_min, settings", "PK + 2 FK + text + int + jsonb", "settings: shuffle, fullscreen…"],
            ["questions", "id, exam_id, type, content, media_url, points, order_idx", "PK, FK, enum, text, text, numeric, int", "type: MCQ | essay"],
            ["exam_sessions", "id, exam_id, status, started_at, ends_at, runtime_state", "PK, FK, enum, timestamptz, jsonb", "runtime: timer server"],
            ["submissions", "id, session_id, user_id, answers, submitted_at, score", "PK, FK, FK, jsonb, timestamptz, numeric", "answers: map câu→lựa chọn"],
            ["integrity_events", "id, session_id, user_id, event_type, payload, created_at", "PK, FK, FK, text, jsonb, timestamptz", "payload: chi tiết sự kiện"],
        ],
        "Bảng 3.7. Thuộc tính chính theo từng bảng (rút gọn — đối chiếu migration thực tế)",
    )
    add_figure(
        doc,
        img["h3_01"],
        "Hình 3.2. Sơ đồ quan hệ thực thể (ER) có nhãn cardinality và tên khóa ngoại",
        width_cm=16.2,
    )
    add_p(
        doc,
        "Tóm tắt thực thể ở bảng 3.8 giúp đối chiếu nhanh khóa chính với các liên kết đã nêu ở bảng 3.6.",
    )
    add_table(
        doc,
        ["Thực thể", "Khóa chính", "Liên kết chính ra ngoài"],
        [
            ["users", "id", "← exams.author_id; ← submissions.user_id; ← integrity_events.user_id"],
            ["subjects", "id", "← exams.subject_id"],
            ["exams", "id", "← questions.exam_id; ← exam_sessions.exam_id"],
            ["questions", "id", "(thuộc exams)"],
            ["exam_sessions", "id", "← submissions.session_id; ← integrity_events.session_id"],
            ["submissions", "id", "session_id + user_id (thường unique theo phiên)"],
            ["integrity_events", "id", "session_id + user_id + timestamp"],
        ],
        "Bảng 3.8. Tóm tắt thực thể — khóa chính và hướng liên kết",
    )
    add_p(
        doc,
        "Thiết kế khóa ngoại đảm bảo không tồn tại submission “mồ côi” khi session bị xóa sai cách; "
        "trong triển khai thực tế cần chính sách ON DELETE phù hợp (restrict/cascade) theo nghiệp vụ.",
    )

    add_h(doc, "3.7. Autosave, integrity và phục hồi runtime", 2)
    add_p(
        doc,
        "Autosave giảm rủi ro mất bài khi sinh viên gặp sự cố mạng hoặc đóng nhầm tab. Contract "
        "autosave trong dự án quy định payload, tần suất khuyến nghị, và cách máy chủ ghi đè an "
        "toàn. Client nên backoff khi nhận 429/503 để tránh bão request.",
    )
    add_p(
        doc,
        "Integrity events ghi nhận hành vi như thoát fullscreen, chuyển tab, copy/paste nếu được "
        "instrumentation. Giám thị hiển thị timeline để quyết định can thiệp. Cần phân biệt “cảnh "
        "báo” và “vi phạm nặng” theo quy chế thi của đơn vị đào tạo.",
    )
    add_p(
        doc,
        "Phục hồi runtime sau restart máy chủ dựa trên bảng trạng thái exam_runtime_state (khái "
        "niệm) giúp timer không bị “tăng thêm” cho thí sinh một cách bất công. Đồ án nên minh họa "
        "bằng sơ đồ trạng thái: chưa bắt đầu → đang làm → đã nộp / hết giờ → đã chấm.",
    )

    add_h(doc, "3.8. Thiết kế giám sát thi và sự kiện Socket", 2)
    add_p(
        doc,
        "Phòng thi số yêu cầu kênh tín hiệu từ giám thị tới thí sinh: bắt đầu đồng loạt, thông báo "
        "chung, cảnh báo cá nhân, dừng khẩn cấp. Socket.IO room theo examId giúp định tuyến sự kiện "
        "hiệu quả. Server cần log socket id kết hợp user id để truy vết.",
    )
    add_p(
        doc,
        "Khi tích hợp force-submit, backend phải đảm bảo idempotent: nếu thí sinh đã nộp, lệnh "
        "lặp lại không làm hỏng dữ liệu. Đồng thời phải thông báo rõ trên UI để tránh tranh chấp.",
    )


def chuong4(doc: Document) -> None:
    add_h(doc, "CHƯƠNG 4: TRIỂN KHAI VÀ KIỂM THỬ HỆ THỐNG", 1)
    add_h(doc, "4.1. Môi trường phát triển và công cụ", 2)
    add_table(
        doc,
        ["Thành phần", "Phiên bản / ghi chú"],
        [
            ["Node.js", "Khuyến nghị LTS phù hợp với package.json dự án"],
            ["npm", "Quản lý dependency"],
            ["PostgreSQL", "Local hoặc Neon cloud"],
            ["Git", "Quản lý phiên bản mã nguồn"],
            ["Trình duyệt", "Chrome/Edge để kiểm thử fullscreen và DevTools"],
        ],
        "Bảng 4.1. Môi trường phát triển",
    )

    add_h(doc, "4.2. Cấu trúc thư mục dự án", 2)
    add_p(
        doc,
        "GraduationProject/BackEnd/server: mã nguồn API, migrations, scripts SQL, openapi.yaml, "
        "API.md. GraduationProject/FrontEnd/client: ứng dụng React, cấu hình Vite, biến môi trường "
        "theo môi trường dev/prod.",
    )

    add_h(doc, "4.3. Triển khai chức năng theo vai trò", 2)
    add_h(doc, "4.3.1. Sinh viên", 3)
    add_bullets(
        doc,
        [
            "Đăng nhập, chọn đề, làm bài với timer và autosave.",
            "Gửi sự kiện integrity khi vi phạm quy tắc (nếu bật).",
            "Nộp bài, xem điểm và phân tích câu đúng/sai khi hệ thống cho phép.",
        ],
    )
    add_h(doc, "4.3.2. Giáo viên", 3)
    add_bullets(
        doc,
        [
            "Tạo/chỉnh sửa đề (theo quyền API), import Word.",
            "Khởi động runtime, theo dõi Socket, chấm tự luận.",
        ],
    )
    add_h(doc, "4.3.3. Quản trị viên", 3)
    add_bullets(
        doc,
        [
            "Quản lý người dùng, reset mật khẩu, audit log, báo cáo hệ thống.",
            "Một số route tạo/sửa đề chỉ admin trên FE — cần lưu ý khi demo.",
        ],
    )

    add_h(doc, "4.4. Kiểm thử đơn vị và kiểm thử hệ thống", 2)
    add_table(
        doc,
        ["Kịch bản", "Bước kiểm tra", "Kết quả mong đợi"],
        [
            ["Health API", "GET /", "200, thông tin service"],
            ["Đăng nhập SV", "POST /v1/auth/login", "JWT hợp lệ"],
            ["Làm bài + autosave", "Quan sát network", "200, không mất tiến độ"],
            ["Hết giờ", "Chờ timer = 0", "Khóa nộp, trạng thái đúng"],
            ["Socket runtime", "Hai trình duyệt", "Sinh viên nhận tín hiệu"],
        ],
        "Bảng 4.2. Kiểm thử thủ công gợi ý",
    )
    add_p(
        doc,
        "Kiểm thử tự động: chạy npm test trong BackEnd/server (Vitest). Nên mở rộng test tích hợp "
        "và CI trước khi triển khai production.",
    )

    add_h(doc, "4.5. Triển khai môi trường thử nghiệm và production", 2)
    add_p(
        doc,
        "Môi trường cục bộ: backend .env sao chép từ .env.example, điền DATABASE_URL, JWT_SECRET, "
        "cấu hình SMTP nếu dùng reset mật khẩu, và biến AI nếu bật MiniMax. Chạy npm install, "
        "npm run migrate, npm run dev. Frontend npm install, npm run dev, trỏ VITE_API_URL tới "
        "localhost trong .env.development.",
    )
    add_p(
        doc,
        "Môi trường production: bắt buộc HTTPS, biến bí mật lưu trên nền tảng triển khai (Render/"
        "VPS), bật CORS đúng origin FE, gzip/brotli nếu reverse proxy hỗ trợ, và healthcheck cho "
        "process manager. Nên tách DB khỏi app server và bật TLS cho kết nối Postgres.",
    )
    add_p(
        doc,
        "Pipeline CI/CD gợi ý: lint + typecheck + unit test trên PR; build docker image hoặc "
        "artifact; deploy blue/green hoặc rolling nếu có nhiều instance. Ghi log tập trung "
        "(JSON lines) giúp truy vết requestId.",
    )

    add_h(doc, "4.6. Hiệu năng, log và xử lý sự cố", 2)
    add_p(
        doc,
        "Hiệu năng làm bài: tránh re-render toàn màn hình khi autosave; debounce input; virtualize "
        "danh sách câu dài. Hiệu năng server: pooling kết nối DB, tránh N+1 query khi tải đề, cache "
        "đề đã tải trong phiên hợp lý.",
    )
    add_p(
        doc,
        "Khi sự cố: kiểm tra log lỗi validation Joi, mã 401 do token hết hạn, 403 do role, 409 "
        "khi xung đột trạng thái phiên. Với Socket, kiểm tra mất kết nối mạng và cơ chế reconnect "
        "của client thư viện.",
    )


def chuong5_ketluan(doc: Document) -> None:
    add_h(doc, "CHƯƠNG 5: ĐÁNH GIÁ KẾT QUẢ VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_h(doc, "5.1. Kết quả đạt được", 2)
    add_bullets(
        doc,
        [
            "Hoàn thành luồng nghiệp vụ thi trực tuyến cốt lõi trên kiến trúc client–server hiện đại.",
            "Có tài liệu tổng hợp DO_AN_MASTER.md, OpenAPI, mô tả API và phân quyền.",
            "Hỗ trợ đa ngôn ngữ, import Word, Socket.IO cho giám sát.",
        ],
    )
    add_h(doc, "5.2. Hạn chế", 2)
    add_bullets(
        doc,
        [
            "Proctoring mức “mềm”, phụ thuộc trình duyệt.",
            "CI và độ bao phủ test cần tăng.",
            "Socket cần chiến lược scale khi multi-instance.",
        ],
    )
    add_h(doc, "5.3. Hướng phát triển", 2)
    add_bullets(
        doc,
        [
            "Redis adapter cho Socket.IO, quan sát metrics (APM).",
            "Chuẩn hóa contract FE–BE bằng test khớp OpenAPI.",
            "Bổ sung xác thực đa yếu tố, rate limit, WAF.",
        ],
    )

    doc.add_page_break()
    add_h(doc, "PHỤ LỤC: DANH MỤC KIỂM THỬ CHI TIẾT (GỢI Ý)", 1)
    add_p(
        doc,
        "Phụ lục liệt kê các testcase thủ công để sinh viên đánh dấu Pass/Fail và chụp màn hình "
        "minh chứng khi nộp báo cáo. Có thể sao chép sang bảng Excel riêng để quản lý tiến độ.",
    )
    test_cases = [
        "Cài đặt dependency backend thành công, không lỗi peer dependency nghiêm trọng.",
        "Cài đặt dependency frontend thành công, chạy dev server tại cổng mặc định.",
        "Sao chép .env.example → .env và điền DATABASE_URL hợp lệ.",
        "Chạy migration trên database trống, không lỗi SQL.",
        "GET / trả về health hợp lệ.",
        "Truy cập /docs (nếu bật Swagger) hiển thị danh sách endpoint.",
        "Đăng ký người dùng (nếu route public) hoặc tạo user qua admin theo seed.",
        "Đăng nhập admin thành công, token lưu đúng cơ chế client.",
        "Đăng nhập giáo viên thành công, phân quyền route hợp lệ.",
        "Đăng nhập sinh viên thành công.",
        "Sai mật khẩu trả về thông báo lỗi rõ ràng, không lộ stack trace production.",
        "Token hết hạn bị từ chối 401 khi gọi API bảo vệ.",
        "Admin xem danh sách người dùng /v1/users.",
        "Admin khóa/mở khóa hoặc cập nhật user (nếu có) không làm hỏng ràng buộc DB.",
        "Admin truy cập audit log và thấy bản ghi sự kiện quan trọng.",
        "Giáo viên tạo đề thi mới với thời lượng và điểm tối đa hợp lệ.",
        "Giáo viên thêm câu trắc nghiệm 4 phương án, một đáp án đúng.",
        "Giáo viên thêm câu tự luận có rubric/điểm tối đa.",
        "Import Word preview hiển thị cấu trúc câu hỏi không vỡ layout.",
        "Import Word commit lưu DB đúng số câu so với preview.",
        "Ngân hàng câu hỏi: tạo câu và gắn môn học (nếu luồng có).",
        "Sinh viên xem danh sách đề được phép làm.",
        "Sinh viên không thể mở đề của lớp khác nếu không được gán.",
        "Bắt đầu làm bài: timer hiển thị đúng thời lượng cấu hình.",
        "Refresh trang giữa chừng: khôi phục runtime, không cộng thêm thời gian bất thường.",
        "Autosave gửi định kỳ, server trả 200, dữ liệu không mất khi reload.",
        "Mất mạng ngắn: client backoff, không spam server.",
        "Integrity: thoát fullscreen ghi nhận sự kiện (nếu bật).",
        "Integrity: chuyển tab ghi nhận (nếu bật).",
        "Giám thị mở trang proctoring, thấy danh sách thí sinh online.",
        "Socket: giáo viên start-runtime, sinh viên nhận tín hiệu bắt đầu.",
        "Socket: broadcast cảnh báo hiển thị trên máy sinh viên.",
        "Socket: force-submit kết thúc phiên, sinh viên không chỉnh sửa thêm.",
        "Nộp bài thủ công trước giờ: khóa làm bài, trạng thái submitted.",
        "Hết giờ tự động nộp: điểm trắc nghiệm được tính.",
        "Xem kết quả: hiển thị điểm tổng và chi tiết đúng/sai (theo cấu hình đề).",
        "Chấm tự luận: giáo viên nhập điểm, tổng điểm cập nhật đúng.",
        "Export kết quả CSV/Excel (nếu có) mở được bằng Excel.",
        "Shuffle câu theo chương: thứ tự khác nhau giữa các sinh viên (nếu bật).",
        "Đa ngôn ngữ: chuyển vi/en/ja, nhãn UI không vỡ layout.",
        "Đổi mật khẩu người dùng (nếu có) và đăng nhập lại thành công.",
        "Reset mật khẩu qua email (nếu SMTP cấu hình) gửi mail hợp lệ.",
        "Dự đoán điểm AI: khi không có API key thì hệ thống degrade êm, không crash.",
        "Dự đoán điểm AI: khi có key, trả về kết quả trong thời gian chấp nhận được.",
        "Phân tích điểm admin: hiển thị thống kê không lỗi với dataset nhỏ.",
        "System report: không lộ thông tin nhạy cảm trong response.",
        "Chạy npm test backend, toàn bộ test pass trên máy dev.",
        "Chạy npm test frontend (nếu có), không fail blocking.",
        "Build production frontend không lỗi TypeScript.",
        "CORS: FE production chỉ gọi được API cho phép.",
        "Bảo mật: không commit file .env vào git (kiểm tra .gitignore).",
        "Log server: lỗi 500 có stack ở dev, ẩn bớt ở prod (nếu cấu hình).",
        "Tải file media câu hỏi (nếu có) hiển thị đúng MIME.",
        "Đăng xuất: token không còn dùng được cho request tiếp theo (nếu revoke client-side).",
        "Kiểm tra song song 20 tab autosave giả lập: server ổn định (stress nhẹ).",
        "Kiểm tra quyền student không gọi được PATCH grade của giáo viên.",
        "Kiểm tra quyền teacher không truy cập được route chỉ admin trên FE.",
        "Backup DB: export và import lại schema + dữ liệu mẫu thành công.",
        "Boundary: đề 1 phút — timer kết thúc đúng, không âm.",
        "Boundary: đề có 200 câu — scroll và autosave vẫn mượt trên máy tầm trung.",
        "Unicode: tiêu đề đề thi có dấu Tiếng Việt đầy đủ.",
        "Hai tab cùng user: kiểm tra hành vi khóa hoặc cảnh báo theo thiết kế.",
        "Đổi ngôn ngữ giữa chừng khi đang làm bài: nhãn cập nhật, không mất câu trả lời.",
        "Giả lập mất mạng 60s: UI hiển thị offline, autosave retry hợp lý.",
        "Chấm điểm vượt quá điểm tối đa câu: server từ chối validation.",
        "Xóa cookie token: request tiếp theo bị 401 và redirect login.",
        "Truy cập API bằng curl không có token: 401 thống nhất.",
        "Kiểm tra CORS preflight OPTIONS từ origin FE hợp lệ.",
        "Kiểm tra log không in password hoặc token đầy đủ.",
        "Kiểm tra rate limit đăng nhập (nếu triển khai): sau N lần sai bị chặn tạm thời.",
        "Kiểm tra migration trên DB đã có dữ liệu: không làm mất bảng quan trọng.",
        "Kiểm tra export có BOM UTF-8 mở đúng trên Excel Windows.",
        "Kiểm tra quyền teacher không xóa user admin.",
        "Kiểm tra đồng hồ khi server restart giữa ca thi (theo tính năng restore).",
        "Kiểm tra file đính kèm câu hỏi quá lớn bị từ chối 413/400.",
        "Kiểm tra song song 50 user đọc dashboard (smoke nhẹ, ghi nhận kết quả).",
    ]
    for i, line in enumerate(test_cases, 1):
        add_p(doc, f"TC-{i:02d}. {line}", justify=False)

    doc.add_page_break()
    add_h(doc, "KẾT LUẬN", 1)
    add_p(
        doc,
        "Đồ án đã phân tích, thiết kế và triển khai hệ thống thi trực tuyến với các module chính "
        "phù hợp thực tiễn giáo dục số. Hệ thống thể hiện được vai trò của kiến trúc REST, cơ sở "
        "dữ liệu quan hệ, và kênh thời gian thực trong tổ chức thi.",
    )
    add_p(
        doc,
        "Hướng tiếp theo là củng cố bảo mật vận hành, mở rộng kiểm thử tự động và hoàn thiện tài "
        "liệu triển khai production. Sinh viên cần bổ sung hình ảnh minh họa theo danh mục hình ảnh "
        "và chỉnh sửa bìa theo quy định của nhà trường.",
    )

    add_h(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Tài liệu nội bộ dự án: DO_AN_MASTER.md, README.md, BackEnd/server/API.md, openapi.yaml.",
        "BackEnd/server/docs/ROLES_AND_PERMISSIONS.md.",
        "BackEnd/server/EXAM_INTEGRITY_AUTOSAVE_CONTRACT.md, SOCKET_IO_POC.md.",
        "Express.js Documentation — https://expressjs.com/",
        "React Documentation — https://react.dev/",
        "PostgreSQL Documentation — https://www.postgresql.org/docs/",
        "Socket.IO Documentation — https://socket.io/docs/",
        "JWT RFC 7519 — https://www.rfc-editor.org/rfc/rfc7519",
    ]
    for i, r in enumerate(refs, 1):
        add_p(doc, f"[{i}] {r}", justify=False)


def noi_dung_bo_sung_dai(doc: Document) -> None:
    """Thêm mô tả dài để đạt khối lượng báo cáo gần ~20 trang khi in A4, Times New Roman 13, giãn dòng 1.5."""
    add_h(
        doc,
        "NỘI DUNG BỔ SUNG: MÔ TẢ CHI TIẾT LUỒNG NGHIỆP VỤ, VẬN HÀNH VÀ BẢO MẬT",
        1,
    )
    add_p(
        doc,
        "Phần này bổ sung mô tả theo hướng “walkthrough” để sinh viên dễ chuyển thành sơ đồ tuần tự "
        "và ảnh chụp màn hình. Khi chỉnh sửa báo cáo cho đúng quy định khoa, có thể gộp các mục vào "
        "Chương 3 và Chương 4 tương ứng.",
    )
    paras: list[str] = [
        (
            "Bước khởi tạo hệ thống: triển khai PostgreSQL, tạo database trống, cấu hình user DB "
            "ít quyền nhất có thể cho ứng dụng, bật kết nối TLS nếu DB nằm ngoài máy chủ ứng dụng. "
            "Chạy migration theo đúng thứ tự, kiểm tra bảng quan trọng đã tồn tại bằng truy vấn "
            "catalog hoặc công cụ GUI. Khởi động backend ở chế độ dev, xác nhận log không có lỗi "
            "kết nối pool."
        ),
        (
            "Bước tạo dữ liệu mẫu: dùng script seed hoặc SQL insert để tạo admin, giáo viên, sinh "
            "viên, môn học và lớp. Kiểm tra ràng buộc unique trên email để tránh trùng. Ghi lại "
            "tài khoản demo phục vụ hội đồng chấm, đồng thời đổi mật khẩu mặc định trước khi public."
        ),
        (
            "Luồng đăng nhập: client gửi thông tin đăng nhập qua HTTPS, server xác minh bcrypt, phát "
            "JWT có exp. Client lưu token theo policy, gắn Authorization: Bearer cho mọi request "
            "REST. Khi token gần hết hạn, có thể hiện thông báo “phiên sắp hết hạn” để người dùng "
            "lưu bài và đăng nhập lại — tính năng refresh token là hướng mở rộng."
        ),
        (
            "Luồng tạo đề thi: giáo viên nhập metadata đề (thời gian, điểm tối đa, cấu hình shuffle, "
            "yêu cầu fullscreen). Server validate Joi, ghi DB, trả examId. Client chuyển sang màn "
            "soạn câu hỏi: thêm từng câu hoặc import Word. Mỗi lần lưu, server đảm bảo tính nhất quán "
            "giữa exam và questions."
        ),
        (
            "Luồng import Word: upload multipart, giới hạn kích thước, server parse bằng Mammoth, "
            "trả preview HTML/JSON cho giáo viên duyệt. Khi commit, server map sang schema câu hỏi "
            "nội bộ, ghi log nếu có cảnh báo định dạng. Nếu parse lỗi, trả mã 400 kèm thông điệp "
            "thân thiện để người dùng chỉnh file nguồn."
        ),
        (
            "Luồng tổ chức phiên: giáo viên tạo session runtime, hệ thống gán trạng thái “chưa bắt "
            "đầu”. Khi đến giờ, giáo viên bấm bắt đầu, server phát sự kiện Socket tới room exam. "
            "Sinh viên đang ở màn chờ chuyển sang màn làm bài, timer đồng bộ theo server time."
        ),
        (
            "Luồng làm bài chi tiết: client render danh sách câu, lưu map câu→đáp án đã chọn trong "
            "state. Mỗi N giây, client gửi autosave; server merge an toàn, tránh ghi đè nếu phiên "
            "đã khóa. Nếu server trả 409, client hiển thị “phiên đã kết thúc” và chuyển sang màn "
            "kết quả."
        ),
        (
            "Luồng integrity: hook sự kiện fullscreen, visibilitychange, copy/paste (tuỳ cấu hình). "
            "Khi phát hiện, client gửi payload gồm loại sự kiện, timestamp client, metadata nhẹ. "
            "Server lưu audit timeline, giám thị có thể lọc theo mức độ. Cần tránh thu thập dữ liệu "
            "nhạy cảm vượt quá phạm vi đồ án."
        ),
        (
            "Luồng nộp bài: sinh viên xác nhận hộp thoại, client gọi submit, server transaction: "
            "khóa submission, tính điểm trắc nghiệm, ghi trạng thái, trả kết quả tóm tắt. Nếu mạng "
            "lỗi, client retry idempotent theo requestId (hướng mở rộng)."
        ),
        (
            "Luồng chấm tự luận: giáo viên mở grading theo sessionId, xem nội dung, nhập điểm từng "
            "câu hoặc tổng, lưu PATCH. Server kiểm tra quyền teacher/admin, cập nhật điểm và log "
            "thay đổi. Sau chấm, sinh viên xem được phản hồi theo policy đề."
        ),
        (
            "Luồng giám sát: giáo viên/admin mở proctoring, server stream danh sách thí sinh online "
            "qua polling hoặc socket events. Khi phát hiện vi phạm lặp lại, giám thị có thể gửi "
            "cảnh báo chung hoặc riêng. Force-submit chỉ dùng khi có căn cứ quy chế."
        ),
        (
            "Luồng quản trị người dùng: admin CRUD user, reset password, gán role. Mọi thao tác nhạy "
            "cảm nên được audit: ai, làm gì, lúc nào, đối tượng nào. Export audit ra CSV phục vụ "
            "kiểm tra nội bộ."
        ),
        (
            "Luồng thống kê điểm: truy vấn aggregate theo exam/class, hiển thị biểu đồ (nếu FE có). "
            "Cần index và giới hạn khung thời gian truy vấn để tránh full scan bảng lớn."
        ),
        (
            "Luồng đa ngôn ngữ: i18next load bundle theo locale, lưu lựa chọn người dùng ở localStorage "
            "hoặc profile. QA cần rà soát chuỗi dài không tràn layout, đặc biệt nút và bảng."
        ),
        (
            "Bảo mật ứng dụng web: chống SQL injection bằng truy vấn tham số hóa; chống XSS nhờ React "
            "escape mặc định và tránh dangerouslySetInnerHTML; chống CSRF cho cookie-based auth "
            "(nếu dùng) bằng SameSite và CSRF token — đồ án dùng Bearer nên CSRF thấp hơn nhưng vẫn "
            "cần chú ý XSS lấy token."
        ),
        (
            "Bảo mật vận hành: rotate JWT_SECRET, giới hạn quyền file system của process, chạy "
            "reverse proxy nginx, bật rate limit đăng nhập, log failed attempts, cấu hình firewall "
            "chỉ mở cổng 443/80."
        ),
        (
            "Quan sát vận hành: gắn requestId middleware, log JSON, tích hợp APM (OpenTelemetry) là "
            "hướng phát triển. Với Socket, log connect/disconnect và room join để debug mất tín hiệu."
        ),
        (
            "Sao lưu và DR: pg_dump định kỳ, lưu offsite, kiểm tra restore định kỳ. Với Neon/cloud, "
            "dùng snapshot nhà cung cấp; ghi rõ RPO/RTO trong báo cáo triển khai thực tế."
        ),
        (
            "Kiểm thử hiệu năng nhẹ: dùng k6 hoặc Apache Bench gửi song song autosave (ở môi trường "
            "staging) để tìm ngưỡng. Kết quả ghi nhận p95 latency và throughput."
        ),
        (
            "Triển khai container (tuỳ chọn): Dockerfile cho backend, build static FE phục vụ qua "
            "nginx. Healthcheck docker giúp orchestrator restart khi treo."
        ),
        (
            "Pháp lý và đạo đức: thu thập log hành vi thi phải có thông báo và đồng ý (consent) theo "
            "quy định đơn vị; đồ án học thuật nên nêu giới hạn và đề xuất văn bản thông báo mẫu."
        ),
        (
            "Kế hoạch bảo trì: semver cho API, changelog, migration an toàn backward compatible khi "
            "có thể; feature flags cho AI prediction để tắt nhanh khi nhà cung cấp lỗi."
        ),
        (
            "Đào tạo người dùng cuối: video ngắn cho giáo viên về import Word và giám sát; hướng dẫn "
            "sinh viên về fullscreen và autosave; FAQ lỗi mạng."
        ),
        (
            "Checklist nộp đồ án: mã nguồn build được, có README, có tài khoản demo, có file Word "
            "báo cáo kèm hình minh họa theo danh mục, có phụ lục testcase đánh dấu Pass/Fail."
        ),
        (
            "Mở rộng học thuật: so sánh thuật toán chấm tự động partial scoring; nghiên cứu IRT cho "
            "đề thích nghi; đây là hướng độc lập ngoài phạm vi hiện tại nhưng có thể đặt làm hướng "
            "phát triển luận văn sau."
        ),
        (
            "Kết luận phần bổ sung: các đoạn trên nhằm “lấp đầy” chi tiết diễn giải để báo cáo đạt "
            "độ dày tương đương đồ án mẫu 15–25 trang khi in; sinh viên nên rút gọn trùng lặp và "
            "thay bằng hình ảnh/sơ đồ khi hoàn thiện."
        ),
        (
            "Phần tiếp theo mô tả kiểm thử biên (boundary): đề có thời gian 1 phút, 5 phút, 180 phút; "
            "số câu 1, 50, 200; điểm tối đa lẻ; file Word rất nhỏ và rất lớn; token JWT sắp hết hạn "
            "trong lúc làm bài; mạng offline 30 giây rồi online lại. Mục tiêu là phát hiện lỗi logic "
            "timer/autosave và race condition khi nộp bài đúng thời điểm biên."
        ),
        (
            "Kiểm thử tương thích trình duyệt: Chrome, Edge, Firefox (phiên bản gần nhất) với các "
            "chính sách fullscreen khác nhau. Ghi nhận khác biệt Fullscreen API và chỉnh UI fallback "
            "khi trình duyệt không hỗ trợ đầy đủ."
        ),
        (
            "Kiểm thử bảo mật nhẹ: thử gọi API admin bằng token student (mong đợi 403), thử chỉnh "
            "sửa payload submit của người khác (mong đợi 403/404), thử replay request autosave cũ "
            "sau khi đã nộp (mong đợi từ chối)."
        ),
        (
            "Kiểm thử hiện tượng “double submit”: double click nút nộp, hoặc gửi hai request song "
            "song; server phải đảm bảo chỉ một trạng thái cuối hợp lệ và không trừ điểm hai lần."
        ),
        (
            "Kiểm thử đồng hồ: đổi múi giờ máy client (nếu có trong lab) để xác nhận server là nguồn "
            "sự thật; client chỉ hiển thị offset, không tự cộng thời gian theo clock local sai lệch."
        ),
        (
            "Kiểm thử dữ liệu Unicode: đề/bài làm có ký tự tiếng Việt, ký tự đặc biệt, emoji trong "
            "phần tự luận (nếu cho phép) để đảm bảo encoding UTF-8 end-to-end và lưu DB không lỗi."
        ),
        (
            "Kiểm thử tải tệp media: ảnh PNG/JPEG lớn, audio ngắn; xác nhận CDN hoặc static path "
            "hoạt động, cache header hợp lý, và không làm treo UI khi tải song song nhiều media."
        ),
        (
            "Kiểm thử phân quyền route FE: truy cập trực tiếp URL admin bằng tài khoản student — "
            "mong đợi redirect hoặc trang 403 theo implement; đồng thời xác nhận không lộ dữ liệu "
            "nhạy cảm trong bundle JS."
        ),
        (
            "Kiểm thử log audit: thực hiện chuỗi hành động admin và xác nhận audit ghi đủ các trường: "
            "actor, action, target, ip, user-agent (nếu có), timestamp theo UTC hoặc timezone rõ ràng."
        ),
        (
            "Kiểm thử backup/restore: sau restore, chạy lại migration idempotent (nếu thiết kế), "
            "đảm bảo không duplicate constraint; kiểm tra sequence id (serial) không lệch gây lỗi "
            "insert."
        ),
        (
            "Kiểm thử song ngữ email: template reset password hiển thị đúng locale; kiểm tra spam "
            "folder và SPF/DKIM khi triển khai thật (ngoài phạm vi đồ án nhưng nên liệt kê)."
        ),
        (
            "Kiểm thử OpenAPI: so khớp một số endpoint quan trọng giữa openapi.yaml và code router; "
            "ghi nhận sai lệch để tránh drift tài liệu — đây là nợ kỹ thuật thường gặp."
        ),
        (
            "Kiểm thử tải đồng thời: mở hai tab cùng user làm cùng một session — mong đợi cảnh báo "
            "hoặc khóa session theo thiết kế “một thiết bị một phiên”; nếu chưa có, ghi nhận là rủi ro."
        ),
        (
            "Kiểm thử downgrade AI: tắt biến môi trường MiniMax, UI không crash, API trả thông báo "
            "“tính năng tạm tắt” phù hợp UX."
        ),
        (
            "Kiểm thử export: file CSV mở bằng Excel không lỗi encoding; delimiter và BOM UTF-8 nếu "
            "cần; dữ liệu nhạy cảm phải được che một phần theo policy."
        ),
        (
            "Kiểm thử hiệu năng DB: explain analyze cho truy vấn dashboard nặng; thêm index nếu "
            "cost sequential scan quá cao; ghi lại số liệu trước/sau tối ưu."
        ),
        (
            "Hoàn thiện báo cáo: chèn sơ đồ use case tổng thể, sơ đồ tuần tự submit, ảnh màn hình "
            "theo danh mục hình; đánh số hình/bảng thống nhất; kiểm tra margin theo template khoa."
        ),
    ]
    for block in paras:
        add_p(doc, block)


def build_document() -> Document:
    root = Path(__file__).resolve().parents[1]
    imgs = ensure_diagrams(root)
    doc = Document()
    set_body_style(doc)
    front_matter(doc)
    danh_muc(doc)
    muc_luc(doc)
    chuong1(doc)
    doc.add_page_break()
    chuong2(doc, imgs)
    doc.add_page_break()
    chuong3(doc, imgs)
    doc.add_page_break()
    chuong4(doc)
    doc.add_page_break()
    noi_dung_bo_sung_dai(doc)
    doc.add_page_break()
    chuong5_ketluan(doc)
    return doc


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "BaoCaoDoAn_OnlineExamination.docx"
    build_document().save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
