"""
Append new Teacher sections (4.4.7-4.4.9) and Admin sections (4.5.7-4.5.8)
into the existing BaoCao_Chuong4_GiaoDien_v2.docx.

Steps:
1. Renumber ADMIN figures first (before inserting anything)
2. Insert teacher additions before 4.5 heading
3. Append admin additions at the end
4. Save to new file (v3)
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

INPUT  = r'c:\VS-Code\GraduationProject\docs\BaoCao_Chuong4_GiaoDien_v2.docx'
OUTPUT = r'c:\VS-Code\GraduationProject\docs\BaoCao_Chuong4_GiaoDien_v3.docx'

doc = Document(INPUT)

# ── helpers ──────────────────────────────────────────────────

def make_heading2(doc, text):
    p = doc.add_paragraph(text, style='Heading 2')
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    return p

def make_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p

def make_figure_placeholder(doc, fig_num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f' Hình 4.{fig_num}. {caption}')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p

# ── Step 1: Find the 4.5 heading index ──────────────────────

admin_heading_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('4.5. Triển khai'):
        admin_heading_idx = i
        break

if admin_heading_idx is None:
    raise ValueError("Could not find 4.5 heading")

print(f"Found 4.5 heading at paragraph index {admin_heading_idx}")

# ── Step 2: Renumber ADMIN figures FIRST (shift +4) ──────────
# Only renumber figures that appear AFTER the 4.5 heading
# Current admin: Hình 4.16 → 4.20, 4.17 → 4.21, ..., 4.21 → 4.25

SHIFT = 4
renumber_count = 0

# Process in REVERSE order to avoid double-renaming (4.16→4.20, then 4.20 would match again)
admin_figure_paras = []
for i, p in enumerate(doc.paragraphs):
    if i <= admin_heading_idx:
        continue
    text = p.text.strip()
    m = re.match(r'Hình 4\.(\d+)\.', text)
    if m:
        old_num = int(m.group(1))
        admin_figure_paras.append((p, old_num))

# Reverse to avoid cascading renames
for p, old_num in reversed(admin_figure_paras):
    new_num = old_num + SHIFT
    for run in p.runs:
        if f'4.{old_num}.' in run.text:
            run.text = run.text.replace(f'4.{old_num}.', f'4.{new_num}.')
            renumber_count += 1
            print(f"  Renumbered Hình 4.{old_num} → 4.{new_num}: {run.text.strip()[:60]}")

print(f"Renumbered {renumber_count} admin figure captions")

# ── Step 3: Insert teacher additions BEFORE 4.5 heading ─────

admin_heading_element = doc.paragraphs[admin_heading_idx]._element
body = doc.element.body

teacher_content = [
    # 4.4.7
    ('heading', '4.4.7. Giao diện quản lý danh sách sinh viên'),
    ('body',
        'Giảng viên có thể quản lý danh sách sinh viên thuộc quyền phụ trách thông qua giao diện quản lý sinh viên. '
        'Bảng danh sách hiển thị thông tin: họ tên, username, email, lớp hành chính và trạng thái tài khoản. '
        'Giảng viên có thể tìm kiếm theo từ khóa, chỉnh sửa thông tin cá nhân của sinh viên, '
        'xóa tài khoản và xem mật khẩu tạm thời của sinh viên mới được tạo. '
        'Chức năng gửi email bảng điểm hàng loạt cho phép giảng viên thông báo kết quả '
        'cho toàn bộ hoặc một nhóm sinh viên được chọn.'
    ),
    ('figure', 16, 'Giao diện quản lý danh sách sinh viên'),

    # 4.4.8
    ('heading', '4.4.8. Giao diện bảng điểm và transcript sinh viên'),
    ('body',
        'Tab bảng điểm trong giao diện quản lý sinh viên cho phép giảng viên chọn bài thi, '
        'tìm kiếm sinh viên và xem bảng điểm dạng bảng với thông tin: họ tên, mã đề, điểm số, '
        'trạng thái chấm bài và thời gian nộp bài. '
        'Giảng viên có thể xuất bảng điểm ra file CSV và gửi email thông báo điểm '
        'cho toàn bộ hoặc nhóm sinh viên được chọn.'
    ),
    ('figure', 17, 'Giao diện bảng điểm theo bài thi'),
    ('body',
        'Giảng viên cũng có thể xem bảng điểm tổng hợp (transcript) của từng sinh viên '
        'thông qua modal hiển thị thông tin cá nhân, danh sách các bài thi đã hoàn thành '
        'với điểm số và xếp loại, cùng phần tổng hợp GPA theo thang điểm 4 và thang điểm 10. '
        'Giao diện hỗ trợ in bảng điểm trực tiếp, xuất file HTML hoặc CSV, '
        'và gửi bảng điểm qua email cho sinh viên.'
    ),
    ('figure', 18, 'Giao diện bảng điểm tổng hợp (transcript) sinh viên'),

    # 4.4.9
    ('heading', '4.4.9. Giao diện trang chủ giảng viên'),
    ('body',
        'Trang chủ giảng viên hiển thị bảng tổng quan hoạt động giảng dạy và quản lý thi cử. '
        'Phần thống kê nhanh bao gồm các thẻ số liệu: tổng số sinh viên quản lý, '
        'số bài thi đã tạo, số phiên thi và số bài chờ chấm điểm. '
        'Bên dưới là bảng danh sách sinh viên với chức năng tìm kiếm và phân trang, '
        'cùng phần hoạt động gần đây hiển thị các sự kiện mới nhất như bài thi được nộp, '
        'phiên thi được mở và kết quả chấm bài. '
        'Giảng viên có thể lọc hoạt động theo trạng thái, khoảng thời gian và từ khóa.'
    ),
    ('figure', 19, 'Giao diện trang chủ giảng viên'),
]

new_elements = []
for item in teacher_content:
    if item[0] == 'heading':
        p = make_heading2(doc, item[1])
    elif item[0] == 'body':
        p = make_body(doc, item[1])
    elif item[0] == 'figure':
        p = make_figure_placeholder(doc, item[1], item[2])
    new_elements.append(p._element)

for el in new_elements:
    body.remove(el)
    admin_heading_element.addprevious(el)

print(f"Inserted {len(new_elements)} teacher paragraphs before 4.5 heading")

# ── Step 4: Append admin additions at the end ────────────────

# 4.5.7
make_heading2(doc, '4.5.7. Giao diện báo cáo hệ thống')
make_body(doc,
    'Giao diện báo cáo hệ thống cung cấp cái nhìn tổng quan về toàn bộ hoạt động của hệ thống '
    'thông qua các chỉ số KPI quan trọng. '
    'Phần thống kê tài khoản hiển thị tổng số tài khoản, số sinh viên, giảng viên và quản trị viên. '
    'Phần thống kê thi cử bao gồm: tổng số bài thi, tổng phiên thi, tỷ lệ hoàn thành, '
    'tỷ lệ đạt và điểm trung bình toàn hệ thống với các thanh tiến trình (progress bar) trực quan.'
)
make_body(doc,
    'Phần thống kê toàn vẹn bài thi hiển thị số vi phạm trong 24 giờ qua, '
    'số phiên bị đánh dấu vi phạm và số bài đang chờ chấm điểm thủ công. '
    'Bên dưới là bảng danh sách các bài thi gần đây với thông tin chi tiết về số phiên, '
    'trạng thái và kết quả tổng hợp, giúp quản trị viên nhanh chóng đánh giá '
    'hiệu quả vận hành tổng thể của hệ thống.'
)
make_figure_placeholder(doc, 26, 'Giao diện báo cáo hệ thống')

# 4.5.8
make_heading2(doc, '4.5.8. Giao diện trang chủ quản trị viên')
make_body(doc,
    'Trang chủ quản trị viên hiển thị bảng tổng quan toàn hệ thống với các thẻ số liệu: '
    'tổng số tài khoản, tổng số bài thi, tổng phiên thi và số lớp hành chính. '
    'Bảng danh sách sinh viên cho phép tìm kiếm và lọc theo lớp hành chính, '
    'giúp quản trị viên nhanh chóng tra cứu thông tin sinh viên.'
)
make_body(doc,
    'Phần hoạt động gần đây liệt kê các sự kiện mới nhất trong hệ thống '
    'như đăng nhập, tạo tài khoản, nộp bài thi và chấm điểm. '
    'Quản trị viên có thể lọc hoạt động theo trạng thái, khoảng thời gian và từ khóa. '
    'Từ trang chủ, quản trị viên có thể nhanh chóng chuyển đến các trang quản lý chi tiết.'
)
make_figure_placeholder(doc, 27, 'Giao diện trang chủ quản trị viên')

print("Appended 2 new admin sections at the end")

# ── Step 5: Save ─────────────────────────────────────────────

doc.save(OUTPUT)
print(f"\nSaved to {OUTPUT}")

# ── Verify final structure ───────────────────────────────────

print("\n=== Final document structure ===")
doc2 = Document(OUTPUT)
for p in doc2.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style else ''
    if 'Heading' in style:
        print(f"  {'>>> ' if 'Heading 1' in style else '    '}{text}")
    elif 'Hình' in text:
        print(f"        {text}")
